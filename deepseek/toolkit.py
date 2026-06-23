# DeepSeek CLI v7.5 — Tool Registry (90+ Tools)
# File, Web, Code, System, Math, Utility + PDF, DOCX, Image, Video, APK, OCR
# + Live Search, Live Model Search, MCP Real-Time Data (16 tools)
# + Web Browser Control (10 tools): Navigate, Login, Click, Fill, Snapshot, etc.
# + Selenium Browser (12 tools): Real browser automation via Firefox/Geckodriver
# + OCR Tools (v6.1): ocr_read, ocr_url — Tesseract + EasyOCR fallback
# + Document Tools (v7.5): PPTX (4), XLSX (4), DOCX Edit, CSV (2), Convert (1)
# NO USAGE LIMITS — all tools available at all times
#
# v5.5 ADDED: Pydantic validation layer for tool inputs
# v6.1 ADDED: OCR support (pytesseract + easyocr fallback)
# v7.0 ADDED: Selenium browser automation (12 tools)
# v7.5 ADDED: Full PPTX/XLSX/DOCX Edit/CSV/Document Conversion tools (11 tools)

import os
import sys
import json
import warnings
# Monkey-patch warnings.warn to suppress duckduckgo_search renaming RuntimeWarnings
_orig_warn = warnings.warn
def _custom_warn(message, category=None, stacklevel=1, source=None):
    if isinstance(message, str) and ("duckduckgo_search" in message or "renamed to" in message):
        return
    return _orig_warn(message, category, stacklevel, source)
warnings.warn = _custom_warn

warnings.filterwarnings("ignore", category=RuntimeWarning, message=".*duckduckgo_search.*")
warnings.filterwarnings("ignore", category=RuntimeWarning, message=".*renamed to.*")
import subprocess
import math
import random
import datetime
import time
import shutil
import platform
import zipfile
import struct
import re
import io
import tempfile
import traceback
from typing import Any
from pathlib import Path

# Pydantic validation (graceful fallback if not installed)
try:
    from pydantic import BaseModel, ValidationError, create_model
    # Pydantic v2: FieldInfo moved to pydantic.fields
    # Pydantic v1: use Field (returns FieldInfo internally)
    try:
        from pydantic.fields import FieldInfo
    except ImportError:
        # Pydantic v1 fallback
        try:
            from pydantic import Field as _Field
            class FieldInfo:
                def __init__(self, default=None, description='', **kwargs):
                    self.default = default
                    self.description = description
        except ImportError:
            FieldInfo = None
    PYDANTIC_AVAILABLE = True
except ImportError:
    PYDANTIC_AVAILABLE = False
    FieldInfo = None


class ToolRegistry:
    """Registry for all agent tools with Pydantic validation (v5.5)."""

    def __init__(self, memory=None):
        self.tools: dict[str, dict] = {}
        self._validation_models: dict[str, type] = {}
        self._memory = memory
        self._register_all()
        # Build Pydantic validation models for all registered tools
        if PYDANTIC_AVAILABLE:
            self._build_validation_models()

    def register(self, name: str, description: str, parameters: dict, handler):
        self.tools[name] = {
            'description': description,
            'parameters': parameters,
            'handler': handler,
        }

    def get_openai_tools(self) -> list[dict]:
        """Return tools in OpenAI function-calling format."""
        result = []
        for name, tool in self.tools.items():
            result.append({
                'type': 'function',
                'function': {
                    'name': name,
                    'description': tool['description'],
                    'parameters': tool['parameters'],
                }
            })
        return result

    def get_tool_list(self) -> list[dict]:
        """Return list of all tools with names and descriptions."""
        return [
            {'name': name, 'description': t['description']}
            for name, t in self.tools.items()
        ]

    # ══════════════════════════════════════
    # PYDANTIC VALIDATION LAYER (v5.5)
    # ══════════════════════════════════════

    def _build_validation_models(self):
        """Build dynamic Pydantic models from tool parameter schemas."""
        for name, tool in self.tools.items():
            try:
                model = self._create_validation_model(name, tool['parameters'])
                if model:
                    self._validation_models[name] = model
            except Exception:
                # Don't let validation model building crash registration
                pass

    def _create_validation_model(self, tool_name: str, params_schema: dict) -> type | None:
        """Create a Pydantic model from an OpenAI-style JSON Schema."""
        if not params_schema or 'properties' not in params_schema:
            return None

        properties = params_schema.get('properties', {})
        required = set(params_schema.get('required', []))

        if not properties:
            return None

        # Map JSON Schema types to Python types
        type_map = {
            'string': str,
            'integer': int,
            'number': float,
            'boolean': bool,
            'array': list,
            'object': dict,
        }

        fields = {}
        for prop_name, prop_schema in properties.items():
            json_type = prop_schema.get('type', 'string')
            py_type = type_map.get(json_type, str)
            is_required = prop_name in required
            default = ... if is_required else None
            description = prop_schema.get('description', '')

            fields[prop_name] = (
                py_type,
                FieldInfo(default=default, description=description)
            )

        model_name = f'{tool_name}_Input'
        return create_model(model_name, **fields)

    def validate_args(self, tool_name: str, args: dict) -> tuple[dict, str | None]:
        """
        Validate tool arguments against the Pydantic model.
        Returns (validated_args, error_string_or_None).
        """
        if not PYDANTIC_AVAILABLE:
            return args, None
        if tool_name not in self._validation_models:
            return args, None

        model_cls = self._validation_models[tool_name]
        try:
            validated = model_cls(**args)
            if hasattr(validated, 'model_dump'):
                return validated.model_dump(), None
            else:
                return validated.dict(), None
        except ValidationError as e:
            # Build friendly error message
            errors = []
            for err in e.errors():
                field = err.get('loc', ['?'])[0]
                msg = err.get('msg', 'invalid')
                expected = err.get('type', '')
                errors.append(f"  {field}: {msg} ({expected})")
            error_msg = f"Validation failed for {tool_name}:\n" + "\n".join(errors)
            return args, error_msg
        except Exception as e:
            return args, f"Validation error for {tool_name}: {e}"

    def execute(self, name: str, arguments: dict) -> str:
        if name not in self.tools:
            return f"[ERROR] Unknown tool '{name}'"

        # ── v5.5: Validate arguments before execution ──
        validated_args, validation_error = self.validate_args(name, arguments)
        if validation_error:
            # Return validation error but don't execute
            return f"[ERROR] {validation_error}"

        try:
            result = self.tools[name]['handler'](validated_args)
            # Check for empty/None result
            if result is None:
                return "[WARNING] Tool returned no output"
            return result
        except TypeError as e:
            # Likely wrong argument types — give helpful error
            return f"[ERROR] Wrong arguments for {name}: {e}"
        except KeyError as e:
            # Missing required argument
            return f"[ERROR] Missing argument for {name}: {e}"
        except Exception as e:
            # Full traceback for debugging
            tb = traceback.format_exc()
            return f"[ERROR] {name} failed: {e}\n{tb}"

    def _register_all(self):
        self._register_file_tools()
        self._register_web_tools()
        self._register_code_tools()
        self._register_system_tools()
        self._register_math_tools()
        self._register_utility_tools()
        self._register_pdf_tools()
        self._register_docx_tools()
        self._register_image_tools()
        self._register_ocr_tools()
        self._register_video_tools()
        self._register_apk_tools()
        self._register_search_tools()
        self._register_mcp_tools()
        self._register_browser_tools()
        self._register_selenium_tools()
        self._register_doc_tools()
        self._register_mcp_client_tools()
        self._register_agent_tools()

    def _register_agent_tools(self):
        """Register multi-agent delegation tools."""
        def delegate_to_agent(args):
            profile = args.get('profile', 'general')
            task = args.get('task', '')
            if not task:
                return 'Error: task is required'
            from .multi_agent import multi_agent_manager
            return multi_agent_manager.delegate(profile, task, self)
        self.register(
            'delegate',
            'Delegate a task to a specialized agent. Profiles: general (all-purpose), coder (programming), researcher (web search), filesystem (file ops), reasoner (step-by-step reasoning)',
            {
                'type': 'object',
                'properties': {
                    'profile': {'type': 'string', 'description': 'Agent profile: general, coder, researcher, filesystem, reasoner'},
                    'task': {'type': 'string', 'description': 'The task description for the sub-agent'},
                },
                'required': ['task'],
            },
            delegate_to_agent
        )

        def delegate_concurrent(args):
            tasks = args.get('tasks', [])
            if not tasks or not isinstance(tasks, list):
                # Fallback: single task via profile+task fields
                profile = args.get('profile', 'general')
                single_task = args.get('task', '')
                if single_task:
                    tasks = [{'profile': profile, 'task': single_task}]
            if not tasks:
                return 'Error: tasks list is required. Use format: {"tasks": [{"profile": "...", "task": "..."}]}'
            from .multi_agent import multi_agent_manager
            task_list = [(t.get('profile', 'general'), t.get('task', '')) for t in tasks]
            task_list = [(p, t) for p, t in task_list if t]
            if not task_list:
                return 'Error: no valid tasks provided'
            ids = multi_agent_manager.run_concurrent_async(task_list, self)
            profiles_str = ', '.join(ids)
            return (f'Started {len(ids)} agents in background: {profiles_str}\n'
                    'Use Ctrl+X then Down to view progress.\n'
                    'Left/Right to switch between agents.\n'
                    'Up or Ctrl+X to return to main chat.')
        self.register(
            'delegate_concurrent',
            'Run multiple tasks concurrently using different specialized agents. Each task runs in parallel.',
            {
                'type': 'object',
                'properties': {
                    'tasks': {
                        'type': 'array',
                        'items': {
                            'type': 'object',
                            'properties': {
                                'profile': {'type': 'string', 'description': 'Agent profile: general, coder, researcher, filesystem, reasoner'},
                                'task': {'type': 'string', 'description': 'The task description'},
                            },
                            'required': ['task'],
                        },
                        'description': 'List of tasks to run concurrently',
                    },
                    'profile': {'type': 'string', 'description': 'Agent profile for single task (fallback if tasks not provided)'},
                    'task': {'type': 'string', 'description': 'Single task description (fallback if tasks not provided)'},
                },
                'required': [],
            },
            delegate_concurrent
        )

    # ══════════════════════════════════════
    # MCP REAL-TIME DATA TOOLS (v5.2)
    # Integrated via Model Context Protocol (mcp)
    # 16 tools: datetime, calendar, news, weather, currency,
    # stock, crypto, holidays, timezone, countdown, sun, qibla, etc.
    # ══════════════════════════════════════

    def _register_mcp_tools(self):
        try:
            from .mcp_tools import get_mcp_tool_definitions, execute_mcp_tool

            for defn in get_mcp_tool_definitions():
                fn = defn['function']
                name = fn['name']
                # Capture name in closure
                self.register(
                    name,
                    fn['description'],
                    fn['parameters'],
                    (lambda n: lambda args: execute_mcp_tool(n, args))(name)
                )
        except ImportError:
            pass  # MCP not installed — skip real-time tools
        except Exception:
            pass  # Error loading MCP — skip silently

    # ══════════════════════════════════════
    # MCP CLIENT TOOLS (v7.7)
    # Real MCP protocol: connects to external MCP servers
    # Discovers and registers tools from Canva, Context7, GitHub, etc.
    # ══════════════════════════════════════

    def _register_mcp_client_tools(self):
        """MCP tools registered on-demand via /mcp connect — no auto-connect."""
        pass

    # ══════════════════════════════════════
    # LIVE SEARCH & MODEL SEARCH TOOLS (v5.2)
    # ══════════════════════════════════════

    def _register_search_tools(self):
        self.register(
            'live_search',
            'LIVE web search: searches DuckDuckGo, Google News, and Bing News in real-time. '
            'Returns titles, URLs, snippets, and source info. Use when you need up-to-date information, '
            'news, or current events. More reliable than basic web_search.',
            {
                'type': 'object',
                'properties': {
                    'query': {'type': 'string', 'description': 'Search query'},
                    'max_results': {'type': 'integer', 'description': 'Max results per source (default 5)'},
                    'source': {'type': 'string', 'description': 'Source: all, duckduckgo, google, bing (default: all)'}
                },
                'required': ['query']
            },
            lambda args: self._live_search(args['query'], args.get('max_results', 5), args.get('source', 'all'))
        )

        self.register(
            'search_models',
            'Search available models from the current AI provider in real-time. Fetches the model list '
            'from the provider API and filters by query. Returns model IDs, names, and context info.',
            {
                'type': 'object',
                'properties': {
                    'query': {'type': 'string', 'description': 'Search query to filter models (e.g. "gpt", "llama", "claude")'},
                    'max_results': {'type': 'integer', 'description': 'Max models to return (default 20)'}
                },
                'required': []
            },
            lambda args: self._search_models(args.get('query', ''), args.get('max_results', 20))
        )

    # ══════════════════════════════════════
    # FILE TOOLS
    # ══════════════════════════════════════

    def _register_file_tools(self):
        self.register(
            'read_file',
            'Read contents of a file from filesystem. Use absolute or relative paths.',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'File path to read'}
                },
                'required': ['path']
            },
            lambda args: self._read_file(args['path'])
        )

        self.register(
            'write_file',
            'Write content to a file. Creates parent dirs if needed. Overwrites existing.',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'File path to write'},
                    'content': {'type': 'string', 'description': 'Content to write'}
                },
                'required': ['path', 'content']
            },
            lambda args: self._write_file(args['path'], args['content'])
        )

        self.register(
            'edit_file',
            'Edit an existing file by finding and replacing text. Shows a diff of changes. '
            'Use this instead of write_file for targeted edits.',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'File path to edit'},
                    'old_string': {'type': 'string', 'description': 'Text to find (must exist in file)'},
                    'new_string': {'type': 'string', 'description': 'Text to replace with'}
                },
                'required': ['path', 'old_string', 'new_string']
            },
            lambda args: self._edit_file(args['path'], args['old_string'], args['new_string'])
        )

        self.register(
            'list_files',
            'List files and directories at a given path. Like ls/dir command.',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'Directory path (default: current)'},
                    'pattern': {'type': 'string', 'description': 'Optional glob filter (e.g. *.py)'}
                },
                'required': []
            },
            lambda args: self._list_files(args.get('path', '.'), args.get('pattern', ''))
        )

        self.register(
            'delete_file',
            'Delete a file or empty directory from filesystem.',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'File or directory path to delete'}
                },
                'required': ['path']
            },
            lambda args: self._delete_file(args['path'])
        )

        self.register(
            'file_info',
            'Get detailed information about a file (size, modified, type, etc).',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'File path'}
                },
                'required': ['path']
            },
            lambda args: self._file_info(args['path'])
        )

    # ══════════════════════════════════════
    # WEB TOOLS
    # ══════════════════════════════════════

    def _register_web_tools(self):
        self.register(
            'web_search',
            'Search the web using DuckDuckGo. Returns titles, URLs, and snippets.',
            {
                'type': 'object',
                'properties': {
                    'query': {'type': 'string', 'description': 'Search query'},
                    'max_results': {'type': 'integer', 'description': 'Max results (default 5)'}
                },
                'required': ['query']
            },
            lambda args: self._web_search(args['query'], args.get('max_results', 5))
        )

        self.register(
            'web_fetch',
            'Fetch and extract text content from a URL. For reading web pages.',
            {
                'type': 'object',
                'properties': {
                    'url': {'type': 'string', 'description': 'URL to fetch'}
                },
                'required': ['url']
            },
            lambda args: self._web_fetch(args['url'])
        )

    # ══════════════════════════════════════
    # CODE TOOLS
    # ══════════════════════════════════════

    def _register_code_tools(self):
        self.register(
            'run_code',
            'Execute Python code and return stdout/stderr output. Useful for calculations, data processing.',
            {
                'type': 'object',
                'properties': {
                    'code': {'type': 'string', 'description': 'Python code to execute'},
                    'timeout': {'type': 'integer', 'description': 'Timeout in seconds (default 30)'}
                },
                'required': ['code']
            },
            lambda args: self._run_code(args['code'], args.get('timeout', 30))
        )

        self.register(
            'run_shell',
            'Execute a shell/bash command and return output. Use for system operations.',
            {
                'type': 'object',
                'properties': {
                    'command': {'type': 'string', 'description': 'Shell command to execute'},
                    'timeout': {'type': 'integer', 'description': 'Timeout in seconds (default 30)'}
                },
                'required': ['command']
            },
            lambda args: self._run_shell(args['command'], args.get('timeout', 30))
        )

        self.register(
            'install_package',
            'Install a Python package via pip. Use when code needs a missing library.',
            {
                'type': 'object',
                'properties': {
                    'package': {'type': 'string', 'description': 'Package name to install'}
                },
                'required': ['package']
            },
            lambda args: self._install_package(args['package'])
        )

    # ══════════════════════════════════════
    # SYSTEM TOOLS
    # ══════════════════════════════════════

    def _register_system_tools(self):
        self.register(
            'system_info',
            'Get system information: OS, CPU, memory, disk, Python version.',
            {
                'type': 'object',
                'properties': {},
                'required': []
            },
            lambda args: self._system_info()
        )

        self.register(
            'process_list',
            'List running processes on the system.',
            {
                'type': 'object',
                'properties': {
                    'filter_name': {'type': 'string', 'description': 'Optional process name filter'}
                },
                'required': []
            },
            lambda args: self._process_list(args.get('filter_name', ''))
        )

        self.register(
            'disk_usage',
            'Get disk usage information for filesystem.',
            {
                'type': 'object',
                'properties': {},
                'required': []
            },
            lambda args: self._disk_usage()
        )

        self.register(
            'network_info',
            'Get network information: hostname, IP, interfaces.',
            {
                'type': 'object',
                'properties': {},
                'required': []
            },
            lambda args: self._network_info()
        )

        self.register(
            'env_vars',
            'Get environment variables. Optionally filter by name pattern.',
            {
                'type': 'object',
                'properties': {
                    'filter': {'type': 'string', 'description': 'Optional name filter (e.g. PATH)'}
                },
                'required': []
            },
            lambda args: self._env_vars(args.get('filter', ''))
        )

    # ══════════════════════════════════════
    # MATH TOOLS
    # ══════════════════════════════════════

    def _register_math_tools(self):
        self.register(
            'calculate',
            'Evaluate a mathematical expression. Supports +, -, *, /, ^, sqrt, sin, cos, tan, log, pi, e.',
            {
                'type': 'object',
                'properties': {
                    'expression': {'type': 'string', 'description': 'Math expression to evaluate'}
                },
                'required': ['expression']
            },
            lambda args: self._calculate(args['expression'])
        )

        self.register(
            'unit_convert',
            'Convert between units: length, weight, temperature, data, time.',
            {
                'type': 'object',
                'properties': {
                    'value': {'type': 'number', 'description': 'Value to convert'},
                    'from_unit': {'type': 'string', 'description': 'Source unit'},
                    'to_unit': {'type': 'string', 'description': 'Target unit'}
                },
                'required': ['value', 'from_unit', 'to_unit']
            },
            lambda args: self._unit_convert(args['value'], args['from_unit'], args['to_unit'])
        )

    # ══════════════════════════════════════
    # UTILITY TOOLS
    # ══════════════════════════════════════

    def _register_utility_tools(self):
        self.register(
            'timestamp',
            'Get current date/time. Optionally convert Unix timestamp or format.',
            {
                'type': 'object',
                'properties': {
                    'unix_ts': {'type': 'number', 'description': 'Optional Unix timestamp to convert'},
                    'timezone': {'type': 'string', 'description': 'Timezone (default: local)'}
                },
                'required': []
            },
            lambda args: self._timestamp(args.get('unix_ts'), args.get('timezone', 'local'))
        )

        self.register(
            'text_transform',
            'Transform text: uppercase, lowercase, reverse, word_count, char_count, base64 encode/decode.',
            {
                'type': 'object',
                'properties': {
                    'text': {'type': 'string', 'description': 'Text to transform'},
                    'operation': {'type': 'string', 'description': 'Operation: upper, lower, reverse, word_count, char_count, base64_encode, base64_decode, title, strip, slug'}
                },
                'required': ['text', 'operation']
            },
            lambda args: self._text_transform(args['text'], args['operation'])
        )

        self.register(
            'json_parse',
            'Parse, format, or query JSON data. Can pretty-print or extract fields.',
            {
                'type': 'object',
                'properties': {
                    'json_string': {'type': 'string', 'description': 'JSON string to parse'},
                    'query': {'type': 'string', 'description': 'Optional dot-notation field to extract'},
                    'pretty': {'type': 'boolean', 'description': 'Pretty print (default true)'}
                },
                'required': ['json_string']
            },
            lambda args: self._json_parse(args['json_string'], args.get('query'), args.get('pretty', True))
        )

        self.register(
            'generate_uuid',
            'Generate random UUIDs. Useful for creating unique identifiers.',
            {
                'type': 'object',
                'properties': {
                    'count': {'type': 'integer', 'description': 'Number of UUIDs to generate (default 1)'}
                },
                'required': []
            },
            lambda args: self._generate_uuid(args.get('count', 1))
        )

        self.register(
            'random_number',
            'Generate random numbers within a range. Can generate integers or floats.',
            {
                'type': 'object',
                'properties': {
                    'min_val': {'type': 'number', 'description': 'Minimum value (default 1)'},
                    'max_val': {'type': 'number', 'description': 'Maximum value (default 100)'},
                    'integer': {'type': 'boolean', 'description': 'Integer only (default true)'}
                },
                'required': []
            },
            lambda args: self._random_number(
                args.get('min_val', 1), args.get('max_val', 100),
                args.get('integer', True)
            )
        )

        self.register(
            'base64_tool',
            'Encode or decode Base64 strings and files.',
            {
                'type': 'object',
                'properties': {
                    'data': {'type': 'string', 'description': 'String to encode/decode'},
                    'mode': {'type': 'string', 'description': 'encode or decode'}
                },
                'required': ['data', 'mode']
            },
            lambda args: self._base64_tool(args['data'], args['mode'])
        )

        self.register(
            'regex_test',
            'Test regular expressions against text. Shows matches and groups.',
            {
                'type': 'object',
                'properties': {
                    'pattern': {'type': 'string', 'description': 'Regex pattern'},
                    'text': {'type': 'string', 'description': 'Text to test against'},
                    'flags': {'type': 'string', 'description': 'Flags: i=ignorecase, m=multiline, s=dotall'}
                },
                'required': ['pattern', 'text']
            },
            lambda args: self._regex_test(args['pattern'], args['text'], args.get('flags', ''))
        )

        self.register(
            'sort_data',
            'Sort lines of text. Can sort alphabetically, numerically, reverse, unique.',
            {
                'type': 'object',
                'properties': {
                    'text': {'type': 'string', 'description': 'Text with lines to sort'},
                    'mode': {'type': 'string', 'description': 'alpha, numeric, reverse, unique, length'},
                    'reverse': {'type': 'boolean', 'description': 'Reverse order'}
                },
                'required': ['text']
            },
            lambda args: self._sort_data(args['text'], args.get('mode', 'alpha'), args.get('reverse', False))
        )

        self.register(
            'count_text',
            'Count occurrences: lines, words, characters, specific pattern matches in text.',
            {
                'type': 'object',
                'properties': {
                    'text': {'type': 'string', 'description': 'Text to analyze'},
                    'pattern': {'type': 'string', 'description': 'Optional pattern to count'}
                },
                'required': ['text']
            },
            lambda args: self._count_text(args['text'], args.get('pattern'))
        )

    # ══════════════════════════════════════
    # PDF TOOLS (v5.0)
    # ══════════════════════════════════════

    def _register_pdf_tools(self):
        self.register(
            'read_pdf',
            'Read and extract text from PDF files. Supports page ranges.',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'Path to PDF file'},
                    'pages': {'type': 'string', 'description': 'Page range: "all", "1-5", "1,3,5" (default: all)'}
                },
                'required': ['path']
            },
            lambda args: self._read_pdf(args['path'], args.get('pages', 'all'))
        )

        self.register(
            'create_pdf',
            'Create PDF documents with formatted content. Supports headings, paragraphs, lists, tables.',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'Output PDF path'},
                    'title': {'type': 'string', 'description': 'Document title'},
                    'content': {'type': 'string', 'description': 'Content in markdown-like format: # heading, ## sub, - list, normal text'}
                },
                'required': ['path', 'content']
            },
            lambda args: self._create_pdf(args['path'], args['content'], args.get('title', ''))
        )

        self.register(
            'pdf_edit',
            'Edit PDF files: merge, split, watermark, rotate, extract pages, get info.',
            {
                'type': 'object',
                'properties': {
                    'operation': {'type': 'string', 'description': 'Operation: merge, split, watermark, rotate, extract, info'},
                    'path': {'type': 'string', 'description': 'PDF file path (or first file for merge)'},
                    'path2': {'type': 'string', 'description': 'Second PDF path (for merge)'},
                    'output': {'type': 'string', 'description': 'Output file path'},
                    'text': {'type': 'string', 'description': 'Watermark text (for watermark op)'},
                    'angle': {'type': 'integer', 'description': 'Rotation angle degrees (for rotate op)'},
                    'pages': {'type': 'string', 'description': 'Pages to extract: "1-3", "1,2,5" (for extract op)'}
                },
                'required': ['operation', 'path']
            },
            lambda args: self._pdf_edit(args)
        )

    # ══════════════════════════════════════
    # DOCX TOOLS (v5.0)
    # ══════════════════════════════════════

    def _register_docx_tools(self):
        self.register(
            'read_docx',
            'Read content from Word DOCX files. Extracts headings, paragraphs, lists, and tables.',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'Path to DOCX file'}
                },
                'required': ['path']
            },
            lambda args: self._read_docx(args['path'])
        )

        self.register(
            'create_docx',
            'Create Word DOCX documents with full formatting. Supports headings, lists, tables, bold, italic.',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'Output DOCX path'},
                    'title': {'type': 'string', 'description': 'Document title'},
                    'content': {'type': 'string', 'description': 'Content: # heading, ## sub, ### h3, - bullet, 1. numbered, **bold**, *italic*, normal text'},
                    'tables': {'type': 'string', 'description': 'JSON string of tables: [{"headers":["Col1","Col2"], "rows":[["a","b"]]}]'}
                },
                'required': ['path', 'content']
            },
            lambda args: self._create_docx(args['path'], args['content'], args.get('title', ''), args.get('tables', ''))
        )

        self.register(
            'todowrite',
            'Display and optionally save a todo/task list. Shows [✓]/[ ] items with bordered opencode-style display.',
            {
                'type': 'object',
                'properties': {
                    'title': {'type': 'string', 'description': 'Todo list title (e.g. "# Todos")'},
                    'items': {
                        'type': 'array',
                        'items': {
                            'type': 'object',
                            'properties': {
                                'text': {'type': 'string', 'description': 'Task description'},
                                'done': {'type': 'boolean', 'description': 'Whether task is completed'}
                            },
                            'required': ['text']
                        },
                        'description': 'List of todo items'
                    },
                    'path': {'type': 'string', 'description': 'Optional file path to save as .txt or .md'}
                },
                'required': ['title', 'items']
            },
            lambda args: self._todowrite(args['title'], args['items'], args.get('path', ''))
        )

        self.register(
            'todolist_get',
            'Show the current todo list stored in session memory.',
            {
                'type': 'object',
                'properties': {},
                'required': []
            },
            lambda args: self._todolist_get()
        )

        self.register(
            'todolist_update',
            'Mark a todo item as done or update its text. Use index to specify which item.',
            {
                'type': 'object',
                'properties': {
                    'index': {'type': 'integer', 'description': 'Index of the item to update (0-based)'},
                    'done': {'type': 'boolean', 'description': 'Set to true to mark as completed'},
                    'text': {'type': 'string', 'description': 'Optional new text for the item'}
                },
                'required': ['index', 'done']
            },
            lambda args: self._todolist_update(args['index'], args.get('done'), args.get('text', ''))
        )

        self.register(
            'todolist_update_all',
            'Mark ALL todo items as done at once. Fast way to complete the entire list.',
            {
                'type': 'object',
                'properties': {},
                'required': []
            },
            lambda args: self._todolist_update_all()
        )

        self.register(
            'docx_info',
            'Get detailed information about a DOCX file: metadata, word count, sections, tables.',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'Path to DOCX file'}
                },
                'required': ['path']
            },
            lambda args: self._docx_info(args['path'])
        )

    # ══════════════════════════════════════
    # IMAGE TOOLS (v5.0)
    # ══════════════════════════════════════

    def _register_image_tools(self):
        self.register(
            'image_view',
            'Display image in terminal as ASCII art with metadata. Works with JPG, PNG, GIF, BMP, WebP.',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'Path to image file'},
                    'width': {'type': 'integer', 'description': 'ASCII art width in chars (default 80)'}
                },
                'required': ['path']
            },
            lambda args: self._image_view(args['path'], args.get('width', 80))
        )

        self.register(
            'image_info',
            'Get detailed image metadata: dimensions, format, mode, EXIF data, file size, color info.',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'Path to image file'}
                },
                'required': ['path']
            },
            lambda args: self._image_info(args['path'])
        )

    # ══════════════════════════════════════
    # DOCUMENT TOOLS (v7.5)
    # PPTX: read, create, edit, info
    # XLSX: read, create, edit, info
    # DOCX: edit (append, replace, tables, properties)
    # CSV: read, create
    # Conversion: xlsx<->csv, md->docx/pdf, txt->docx, json->xlsx/csv
    # ══════════════════════════════════════

    def _register_doc_tools(self):
        try:
            from .doc_tools import get_doc_tool_definitions, execute_doc_tool

            for defn in get_doc_tool_definitions():
                fn = defn['function']
                name = fn['name']
                self.register(
                    name,
                    fn['description'],
                    fn['parameters'],
                    (lambda n: lambda args: execute_doc_tool(n, args))(name)
                )
        except ImportError:
            pass
        except Exception:
            pass

    # ══════════════════════════════════════
    # OCR TOOLS (v6.1)
    # pytesseract (primary) + easyocr (fallback)
    # ══════════════════════════════════════

    def _register_ocr_tools(self):
        self.register(
            'ocr_read',
            'Read and extract text from an image file using OCR (Optical Character Recognition). '
            'Supports JPG, PNG, BMP, TIFF, WebP, GIF. Returns the extracted text. '
            'Requires pytesseract or easyocr to be installed. Use for screenshots, scanned documents, '
            'signs, receipts, handwritten text, or any image containing readable text.',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'Path to the image file'},
                    'language': {'type': 'string',
                                'description': 'Language(s) for OCR. Tesseract: "eng", "eng+ind", "jpn", etc. '
                                               'EasyOCR: "en", "id", "ja", etc. Default: "eng"'},
                    'engine': {'type': 'string',
                              'description': 'OCR engine: "tesseract" (fast, needs tesseract-ocr installed), '
                                             '"easyocr" (slower but no system dependency), or "auto" (try tesseract first, fallback to easyocr). '
                                             'Default: "auto"'},
                    'preprocess': {'type': 'string',
                                   'description': 'Image preprocessing: "none" (default), "grayscale", "threshold", '
                                                  '"denoise", "sharpen", "enhance". Improves OCR accuracy on low-quality images.'},
                },
                'required': ['path']
            },
            lambda args: self._ocr_read(
                args['path'],
                args.get('language', 'eng'),
                args.get('engine', 'auto'),
                args.get('preprocess', 'none'),
            )
        )

        self.register(
            'ocr_url',
            'Download an image from a URL and extract text using OCR. Supports the same options as ocr_read. '
            'Useful for reading text from online images, screenshots shared as links, etc.',
            {
                'type': 'object',
                'properties': {
                    'url': {'type': 'string', 'description': 'URL of the image to OCR'},
                    'language': {'type': 'string',
                                'description': 'Language(s) for OCR. Default: "eng"'},
                    'engine': {'type': 'string',
                              'description': 'OCR engine: "tesseract", "easyocr", or "auto". Default: "auto"'},
                    'preprocess': {'type': 'string',
                                   'description': 'Image preprocessing: "none", "grayscale", "threshold", "denoise", "sharpen", "enhance"'},
                },
                'required': ['url']
            },
            lambda args: self._ocr_url(
                args['url'],
                args.get('language', 'eng'),
                args.get('engine', 'auto'),
                args.get('preprocess', 'none'),
            )
        )

    # ══════════════════════════════════════
    # VIDEO TOOLS (v5.0)
    # ══════════════════════════════════════

    def _register_video_tools(self):
        self.register(
            'video_info',
            'Get video metadata: duration, resolution, codec, bitrate, FPS, audio info. Uses ffprobe.',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'Path to video file'}
                },
                'required': ['path']
            },
            lambda args: self._video_info(args['path'])
        )

        self.register(
            'video_play',
            'Open video/image/audio in system player. Uses termux-open (Termux) or xdg-open (Linux).',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'Path to media file'}
                },
                'required': ['path']
            },
            lambda args: self._video_play(args['path'])
        )

    # ══════════════════════════════════════
    # APK TOOLS (v5.0 BETA)
    # ══════════════════════════════════════

    def _register_apk_tools(self):
        self.register(
            'apk_analyze',
            '[BETA] Analyze Android APK files. Extract package name, permissions, activities, DEX info, signing.',
            {
                'type': 'object',
                'properties': {
                    'path': {'type': 'string', 'description': 'Path to APK file'}
                },
                'required': ['path']
            },
            lambda args: self._apk_analyze(args['path'])
        )


    # ══════════════════════════════════════
    # WEB BROWSER CONTROL TOOLS (v5.4)
    # Full HTTP-based browser automation: Navigate, Login, Click, Fill, etc.
    # Uses persistent BrowserSession with cookie management
    # ══════════════════════════════════════

    def _register_browser_tools(self):
        self.register(
            'browser_navigate',
            'WEB BROWSER: Navigate to a URL. Opens a web page and returns title, visible text, '
            'links, forms, and images. Maintains session cookies across navigations. '
            'Use this as the first step when interacting with any website.',
            {
                'type': 'object',
                'properties': {
                    'url': {'type': 'string', 'description': 'Full URL to navigate to (e.g. https://example.com)'}
                },
                'required': ['url']
            },
            lambda args: self._browser_navigate(args['url'])
        )

        self.register(
            'browser_login',
            'WEB BROWSER: Login to a website via form. Auto-detects username/password fields. '
            'Maintains session cookies after login for subsequent requests. '
            'Use after browser_navigate to the login page.',
            {
                'type': 'object',
                'properties': {
                    'url': {'type': 'string', 'description': 'Login page URL'},
                    'username': {'type': 'string', 'description': 'Username or email'},
                    'password': {'type': 'string', 'description': 'Password'},
                    'username_field': {'type': 'string', 'description': 'Name/ID of username field (auto-detect if empty)'},
                    'password_field': {'type': 'string', 'description': 'Name/ID of password field (auto-detect if empty)'}
                },
                'required': ['url', 'username', 'password']
            },
            lambda args: self._browser_login(
                args['url'], args['username'], args['password'],
                args.get('username_field', ''), args.get('password_field', '')
            )
        )

        self.register(
            'browser_click',
            'WEB BROWSER: Click a link or button on the current page by text or CSS selector. '
            'Follows the link or submits the parent form. '
            'Call browser_navigate first to load the page.',
            {
                'type': 'object',
                'properties': {
                    'target': {'type': 'string', 'description': 'Link text, button text, or CSS selector to click'},
                    'by': {'type': 'string', 'description': 'Search method: "text" (match text content) or "css" (CSS selector)'}
                },
                'required': ['target']
            },
            lambda args: self._browser_click(args['target'], args.get('by', 'text'))
        )

        self.register(
            'browser_fill_form',
            'WEB BROWSER: Fill form fields and optionally submit. Navigate to the form page first. '
            'form_data is JSON string: {"field_name": "value", ...}',
            {
                'type': 'object',
                'properties': {
                    'url': {'type': 'string', 'description': 'URL of the page with the form'},
                    'form_data': {'type': 'string', 'description': 'JSON string of form data: {"username": "john", "email": "j@mail.com"}'},
                    'submit': {'type': 'boolean', 'description': 'Submit the form after filling (default true)'},
                    'form_index': {'type': 'integer', 'description': 'Which form if multiple (0-based, default 0)'}
                },
                'required': ['url', 'form_data']
            },
            lambda args: self._browser_fill_form(
                args['url'], args['form_data'],
                args.get('submit', True), args.get('form_index', 0)
            )
        )

        self.register(
            'browser_snapshot',
            'WEB BROWSER: Full page snapshot — returns visible text, all links, all forms, '
            'all images, headings structure, metadata, stylesheets, and scripts. '
            'Comprehensive view of the entire page for analysis.',
            {
                'type': 'object',
                'properties': {},
                'required': []
            },
            lambda args: self._browser_snapshot()
        )

        self.register(
            'browser_extract',
            'WEB BROWSER: Extract content from current page using CSS selector. '
            'Returns matching elements with their text, attributes, links, and images.',
            {
                'type': 'object',
                'properties': {
                    'css_selector': {'type': 'string', 'description': 'CSS selector (e.g. "div.article", "table.data tr", "#main-content")'}
                },
                'required': ['css_selector']
            },
            lambda args: self._browser_extract(args['css_selector'])
        )

        self.register(
            'browser_download',
            'WEB BROWSER: Download a file from URL using session cookies. '
            'Saves to specified path or auto-generates filename from URL.',
            {
                'type': 'object',
                'properties': {
                    'url': {'type': 'string', 'description': 'URL of file to download'},
                    'save_path': {'type': 'string', 'description': 'Local save path (auto-generated from URL if empty)'}
                },
                'required': ['url']
            },
            lambda args: self._browser_download(args['url'], args.get('save_path', ''))
        )

        self.register(
            'browser_action',
            'WEB BROWSER: Perform a general web action. Combines navigate, click, fill, and extract '
            'in one tool. Actions: navigate, click, fill, submit, extract, back.',
            {
                'type': 'object',
                'properties': {
                    'action': {'type': 'string', 'description': 'Action type: navigate, click, fill, submit, extract, back'},
                    'url': {'type': 'string', 'description': 'URL (for navigate/fill actions)'},
                    'target': {'type': 'string', 'description': 'Target element text or CSS selector (for click/extract)'},
                    'data': {'type': 'string', 'description': 'JSON data string (for fill action)'},
                    'css_selector': {'type': 'string', 'description': 'CSS selector (for extract action)'}
                },
                'required': ['action']
            },
            lambda args: self._browser_action(args)
        )

        self.register(
            'browser_screenshot',
            'WEB BROWSER: Generate visual text-based rendering of the current page. '
            'Shows page structure: headings, paragraphs, links, images, forms, '
            'tables, buttons, inputs, lists — like a DOM tree visualization.',
            {
                'type': 'object',
                'properties': {},
                'required': []
            },
            lambda args: self._browser_screenshot()
        )

        self.register(
            'browser_cookies',
            'WEB BROWSER: View or manage session cookies. Shows all cookies for the current session '
            'or clears them. Cookies persist across all browser tool calls.',
            {
                'type': 'object',
                'properties': {
                    'action': {'type': 'string', 'description': 'Action: "view" (show cookies) or "clear" (delete all)'}
                },
                'required': []
            },
            lambda args: self._browser_cookies(args.get('action', 'view'))
        )


    # ═══════════════════════════════════════════════════════════════
    # TOOL IMPLEMENTATIONS
    # ═══════════════════════════════════════════════════════════════

    def _read_file(self, path: str) -> str:
        p = Path(path).expanduser()
        if not p.exists():
            return f"Error: File not found: {path}"
        if not p.is_file():
            return f"Error: Not a file: {path}"
        try:
            content = p.read_text(encoding='utf-8', errors='replace')
            lines = content.split('\n')
            if len(lines) > 500:
                preview = '\n'.join(lines[:250])
                return f"[Showing first 250 of {len(lines)} lines]\n\n{preview}\n\n... ({len(lines) - 250} more lines)"
            return content
        except Exception as e:
            return f"Error reading file: {e}"

    def _write_file(self, path: str, content: str) -> str:
        p = Path(path).expanduser()
        try:
            p.parent.mkdir(parents=True, exist_ok=True)
            old_content = p.read_text(encoding='utf-8', errors='replace') if p.exists() else ''
            p.write_text(content, encoding='utf-8')
            lines = content.count('\n') + 1
            result = f"Written {len(content)} chars ({lines} lines) to {path}"
            if old_content:
                diff = self._format_diff(old_content, content)
                if diff:
                    result = f"[DIFF]{diff}[/DIFF]"
            return result
        except Exception as e:
            return f"Error writing file: {e}"

    def _edit_file(self, path: str, old_string: str, new_string: str) -> str:
        p = Path(path).expanduser()
        if not p.exists():
            return f"Error: File not found: {path}"
        try:
            content = p.read_text(encoding='utf-8', errors='replace')
            if old_string not in content:
                return f"Error: old_string not found in {path}"
            new_content = content.replace(old_string, new_string, 1)
            diff = self._format_diff(content, new_content)
            p.write_text(new_content, encoding='utf-8')
            if diff:
                return f"[DIFF]{diff}[/DIFF]"
            return f"Edited {path}"
        except Exception as e:
            return f"Error editing file: {e}"

    @staticmethod
    def _format_diff(old_text: str, new_text: str, max_lines: int = 20) -> str:
        """Generate a compact line-diff with line numbers. Skips if new file."""
        if not old_text.strip():
            return ''
        old_lines = old_text.split('\n')
        new_lines = new_text.split('\n')
        total_lines = max(len(old_lines), len(new_lines))
        pad = len(str(total_lines))
        context = 1
        changes = []
        for i in range(total_lines):
            a = old_lines[i] if i < len(old_lines) else ''
            b = new_lines[i] if i < len(new_lines) else ''
            if a != b:
                changes.append(i)
        if not changes:
            return ''
        shown = set()
        result = []
        for idx in changes:
            if len(result) >= max_lines:
                remaining = len(changes) - changes.index(idx)
                result.append(f"~   ... {remaining} more changes")
                break
            start = max(0, idx - context)
            end = min(total_lines, idx + context + 1)
            for i in range(start, end):
                if i in shown or len(result) >= max_lines:
                    continue
                shown.add(i)
                a = old_lines[i] if i < len(old_lines) else ''
                b = new_lines[i] if i < len(new_lines) else ''
                ln = str(i + 1).rjust(pad)
                if a != b:
                    result.append(f"-{ln} {a}")
                    result.append(f"+{ln} {b}")
                else:
                    result.append(f" {ln} {a}")
        return '\n'.join(result)

    def _list_files(self, path: str, pattern: str) -> str:
        p = Path(path).expanduser()
        if not p.exists():
            return f"Error: Path not found: {path}"
        if not p.is_dir():
            return f"Error: Not a directory: {path}"
        try:
            items = list(p.iterdir())
            if pattern:
                items = [i for i in items if i.match(pattern)]
            items.sort(key=lambda x: (not x.is_dir(), x.name.lower()))
            lines = [f"  {i.name}/" if i.is_dir() else f"  {i.name}" for i in items]
            if not lines:
                return f"No items found in {path}" + (f" matching '{pattern}'" if pattern else "")
            return f"Contents of {path} ({len(lines)} items):\n" + "\n".join(lines)
        except PermissionError:
            return f"Error: Permission denied: {path}"
        except Exception as e:
            return f"Error: {e}"

    def _delete_file(self, path: str) -> str:
        p = Path(path).expanduser()
        if not p.exists():
            return f"Error: Not found: {path}"
        try:
            if p.is_dir():
                shutil.rmtree(p)
                return f"Deleted directory: {path}"
            else:
                p.unlink()
                return f"Deleted file: {path}"
        except Exception as e:
            return f"Error deleting: {e}"

    def _file_info(self, path: str) -> str:
        p = Path(path).expanduser()
        if not p.exists():
            return f"Error: Not found: {path}"
        stat = p.stat()
        size = stat.st_size
        if size > 1024 * 1024:
            size_str = f"{size / (1024*1024):.2f} MB"
        elif size > 1024:
            size_str = f"{size / 1024:.2f} KB"
        else:
            size_str = f"{size} B"
        mtime = datetime.datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
        ftype = 'directory' if p.is_dir() else 'file'
        return (f"Path: {p}\n"
                f"Type: {ftype}\n"
                f"Size: {size_str}\n"
                f"Modified: {mtime}\n"
                f"Permissions: {oct(stat.st_mode)[-3:]}")

    def _web_search(self, query: str, max_results: int) -> str:
        try:
            try:
                from ddgs import DDGS
            except ImportError:
                from duckduckgo_search import DDGS
            results = []
            with DDGS() as ddgs:
                for r in ddgs.text(query, max_results=max_results):
                    results.append(f"[{r.get('title', '')}]({r.get('href', '')})\n  {r.get('body', '')}")
            if not results:
                return f"No results found for: {query}"
            return f"Search results for '{query}':\n\n" + "\n\n".join(results)
        except ImportError:
            return "Error: duckduckgo-search package not installed. Run: pip install duckduckgo-search"
        except Exception as e:
            return f"Search error: {e}"

    def _web_fetch(self, url: str) -> str:
        """Fetch URL and return clean readable content (HTML stripped)."""
        # Try multiple User-Agent headers — some sites block certain UAs
        ua_list = [
            'Mozilla/5.0 (Linux; Android 14; Pixel 8) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Mobile Safari/537.36',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36',
            'Mozilla/5.0 (compatible; Googlebot/2.1; +http://www.google.com/bot.html)',
        ]

        last_error = ''
        for ua in ua_list:
            try:
                import httpx

                headers = {
                    'User-Agent': ua,
                    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                    'Accept-Language': 'en-US,en;q=0.9',
                }

                with httpx.Client(timeout=20, follow_redirects=True) as client:
                    r = client.get(url, headers=headers)

                # 403/429 — try next UA
                if r.status_code in (403, 429):
                    last_error = f"HTTP {r.status_code}"
                    continue

                r.raise_for_status()
                content_type = r.headers.get('content-type', '')

                # Binary content
                if 'text' not in content_type and 'json' not in content_type:
                    return f"Binary content (type: {content_type}), size: {len(r.content)} bytes"

                text = r.text

                # JSON response — return pretty-printed
                if 'json' in content_type or text.strip().startswith(('{', '[')):
                    try:
                        import json
                        obj = json.loads(text)
                        pretty = json.dumps(obj, indent=2, ensure_ascii=False)
                        return pretty[:8000] + (f"\n\n... (truncated, total {len(pretty)} chars)" if len(pretty) > 8000 else "")
                    except (json.JSONDecodeError, ValueError):
                        pass

                # HTML — parse and extract clean text
                if '<html' in text.lower() or '<body' in text.lower() or '<div' in text.lower():
                    clean = self._html_to_text(text)
                    if clean and len(clean.strip()) > 50:
                        if len(clean) > 8000:
                            clean = clean[:8000] + f"\n\n... (truncated, total {len(clean)} chars)"
                        return clean
                    # Fallback to raw if parser produced nothing useful
                    return text[:5000] + (f"\n\n... (truncated, total {len(text)} chars)" if len(text) > 5000 else "")

                # Plain text
                return text[:8000] + (f"\n\n... (truncated, total {len(text)} chars)" if len(text) > 8000 else "")

            except httpx.TimeoutException:
                return "Fetch error: Request timed out (20s)"
            except httpx.ConnectError:
                return "Fetch error: Connection failed. Check internet."
            except httpx.HTTPStatusError as e:
                last_error = f"HTTP {e.response.status_code}"
                continue
            except Exception as e:
                last_error = str(e)
                continue

        return f"Fetch error: {last_error} (all attempts failed)"

    def _html_to_text(self, html: str) -> str:
        """Convert HTML to clean readable text. Tries bs4, falls back to regex."""
        # Try BeautifulSoup first (best quality)
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, 'html.parser')

            # Remove noise elements
            for tag in soup(['script', 'style', 'noscript', 'iframe', 'svg',
                            'header', 'footer', 'nav', 'aside', 'form',
                            'button', 'input', 'select', 'textarea']):
                tag.decompose()

            # Extract title
            title = ''
            title_tag = soup.find('title')
            if title_tag and title_tag.string:
                title = title_tag.string.strip()

            # Extract meta description
            meta_desc = ''
            meta = soup.find('meta', attrs={'name': 'description'})
            if meta and meta.get('content'):
                meta_desc = meta['content'].strip()

            # Get text from body (or whole doc if no body)
            body = soup.find('body') or soup
            text = body.get_text(separator='\n', strip=True)

            # Clean up: remove excessive blank lines
            import re
            text = re.sub(r'\n{3,}', '\n\n', text)
            text = text.strip()

            # Build result with metadata
            parts = []
            if title:
                parts.append(f"Title: {title}")
            if meta_desc:
                parts.append(f"Description: {meta_desc}")
            if text:
                parts.append(text)

            result = '\n\n'.join(parts)
            return result if result.strip() else ''

        except ImportError:
            pass
        except Exception:
            pass

        # Fallback: regex-based HTML stripping (no bs4 needed)
        import re
        # Remove scripts, styles
        text = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<noscript[^>]*>.*?</noscript>', '', html, flags=re.DOTALL | re.IGNORECASE)
        # Remove HTML tags
        text = re.sub(r'<[^>]+>', ' ', text)
        # Decode common HTML entities
        text = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
        text = text.replace('&quot;', '"').replace('&#39;', "'").replace('&nbsp;', ' ')
        # Clean whitespace
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    def _run_code(self, code: str, timeout: int) -> str:
        try:
            cwd = os.environ.get('DEEPSEEK_ORIGINAL_CWD') or os.getcwd()
            if cwd in ('$PWD', '${PWD}'):
                cwd = os.getcwd()
            result = subprocess.run(
                [sys.executable, '-c', code],
                capture_output=True, text=True, timeout=timeout, cwd=cwd
            )
            output = ''
            if result.stdout:
                output += result.stdout
            if result.stderr:
                output += f"\n[stderr]: {result.stderr}"
            if result.returncode != 0:
                output += f"\n[exit code: {result.returncode}]"
            return output.strip() if output.strip() else "(no output)"
        except subprocess.TimeoutExpired:
            return f"Code timed out after {timeout}s"
        except Exception as e:
            return f"Code execution error: {e}"

    def _run_shell(self, command: str, timeout: int) -> str:
        try:
            if os.geteuid() == 0:
                command = re.sub(r'^\s*sudo\s+', '', command)
            else:
                if re.match(r'^\s*sudo\s', command):
                    if not re.search(r'\s-[in]\s', command):
                        command = re.sub(r'^\s*sudo\s', 'sudo -n ', command)
            cwd = os.environ.get('DEEPSEEK_ORIGINAL_CWD') or os.getcwd()
            # Safety: reject literal $PWD (wrapper bug in old installs)
            if cwd in ('$PWD', '${PWD}'):
                cwd = os.getcwd()
            result = subprocess.run(
                command, shell=True, capture_output=True, text=True, timeout=timeout, cwd=cwd
            )
            output = ''
            if result.stdout:
                output += result.stdout
            if result.stderr:
                output += f"\n[stderr]: {result.stderr}"
            return output.strip() if output.strip() else "(no output)"
        except subprocess.TimeoutExpired:
            return f"Command timed out after {timeout}s"
        except Exception as e:
            return f"Shell error: {e}"

    def _install_package(self, package: str) -> str:
        try:
            result = subprocess.run(
                [sys.executable, '-m', 'pip', 'install', '-q', package],
                capture_output=True, text=True, timeout=120
            )
            if result.returncode == 0:
                return f"Package '{package}' installed successfully."
            return f"Install failed: {result.stderr}"
        except Exception as e:
            return f"Install error: {e}"

    def _system_info(self) -> str:
        uname = platform.uname()
        mem = ''
        try:
            total = os.popen('free -h 2>/dev/null || cat /proc/meminfo 2>/dev/null | head -3').read()
            mem = f"\nMemory:\n{total.strip()}" if total.strip() else ''
        except Exception:
            pass
        return (f"System: {uname.system} {uname.release}\n"
                f"Node: {uname.node}\n"
                f"Machine: {uname.machine}\n"
                f"Processor: {uname.processor}\n"
                f"Python: {sys.version}\n"
                f"CWD: {os.getcwd()}{mem}")

    def _process_list(self, filter_name: str) -> str:
        try:
            cmd = 'ps aux 2>/dev/null || ps -ef 2>/dev/null'
            result = subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=10)
            lines = result.stdout.strip().split('\n')
            if filter_name:
                lines = [l for l in lines if filter_name.lower() in l.lower()]
            if len(lines) > 20:
                lines = lines[:20]
                lines.append(f"... (showing 20 of many, use filter to narrow)")
            return '\n'.join(lines) if lines else "No processes found"
        except Exception as e:
            return f"Error: {e}"

    def _disk_usage(self) -> str:
        try:
            result = subprocess.run('df -h 2>/dev/null', shell=True, capture_output=True, text=True, timeout=10)
            return result.stdout.strip() if result.stdout.strip() else "Disk info not available"
        except Exception as e:
            return f"Error: {e}"

    def _network_info(self) -> str:
        try:
            hostname = platform.node()
            ip = ''
            try:
                s = __import__('socket')
                ip = s.gethostbyname(hostname)
            except Exception:
                pass
            interfaces = ''
            try:
                result = subprocess.run('ifconfig 2>/dev/null || ip addr 2>/dev/null || ipconfig 2>/dev/null',
                                       shell=True, capture_output=True, text=True, timeout=10)
                interfaces = f"\nInterfaces:\n{result.stdout.strip()[:1000]}"
            except Exception:
                pass
            return f"Hostname: {hostname}\nIP: {ip}{interfaces}"
        except Exception as e:
            return f"Error: {e}"

    def _env_vars(self, filter_name: str) -> str:
        env = dict(os.environ)
        if filter_name:
            env = {k: v for k, v in env.items() if filter_name.upper() in k.upper()}
        if not env:
            return f"No env vars matching '{filter_name}'"
        lines = [f"{k}={v}" for k, v in sorted(env.items())]
        return '\n'.join(lines)

    def _calculate(self, expression: str) -> str:
        import ast
        import operator
        
        safe_map = {
            'sqrt': math.sqrt, 'sin': math.sin, 'cos': math.cos,
            'tan': math.tan, 'log': math.log, 'log10': math.log10,
            'pi': math.pi, 'e': math.e, 'abs': abs,
            'ceil': math.ceil, 'floor': math.floor, 'pow': pow,
            'round': round,
        }
        
        # AST-based safe evaluation for math expressions
        def safe_eval(node):
            if isinstance(node, ast.Num):  # <number>
                return node.n
            elif isinstance(node, ast.Constant):  # Python 3.8+ Constant
                if isinstance(node.value, (int, float)):
                    return node.value
                raise ValueError("Only numeric constants allowed")
            elif isinstance(node, ast.BinOp):  # <left> <operator> <right>
                ops = {
                    ast.Add: operator.add,
                    ast.Sub: operator.sub,
                    ast.Mult: operator.mul,
                    ast.Div: operator.truediv,
                    ast.Pow: operator.pow,
                    ast.BitXor: operator.pow,  # Map ^ to ** as well
                    ast.Mod: operator.mod,
                    ast.FloorDiv: operator.floordiv
                }
                op = type(node.op)
                if op in ops:
                    return ops[op](safe_eval(node.left), safe_eval(node.right))
                raise ValueError(f"Unsupported operator: {op}")
            elif isinstance(node, ast.UnaryOp):  # <operator> <operand> e.g., -1
                if isinstance(node.op, ast.USub):
                    return -safe_eval(node.operand)
                elif isinstance(node.op, ast.UAdd):
                    return +safe_eval(node.operand)
                raise ValueError(f"Unsupported unary operator: {type(node.op)}")
            elif isinstance(node, ast.Call):  # <func>(<args>)
                if isinstance(node.func, ast.Name):
                    func_name = node.func.id
                    if func_name in safe_map and callable(safe_map[func_name]):
                        args = [safe_eval(arg) for arg in node.args]
                        return safe_map[func_name](*args)
                raise ValueError("Unsupported function call")
            elif isinstance(node, ast.Name):  # Constants like pi, e
                if node.id in safe_map and not callable(safe_map[node.id]):
                    return safe_map[node.id]
                raise ValueError(f"Unsupported variable: {node.id}")
            raise ValueError(f"Unsupported syntax: {type(node)}")

        try:
            # Replace ^ with ** for exponentiation (common LLM behavior)
            expr_clean = expression.replace('^', '**')
            tree = ast.parse(expr_clean, mode='eval')
            result = safe_eval(tree.body)
            return str(result)
        except Exception as e:
            return f"Math error: {e}"

    def _unit_convert(self, value, from_unit: str, to_unit: str) -> str:
        conversions = {
            ('km', 'mi'): lambda v: v * 0.621371,
            ('mi', 'km'): lambda v: v * 1.60934,
            ('kg', 'lb'): lambda v: v * 2.20462,
            ('lb', 'kg'): lambda v: v * 0.453592,
            ('c', 'f'): lambda v: v * 9/5 + 32,
            ('f', 'c'): lambda v: (v - 32) * 5/9,
            ('kb', 'mb'): lambda v: v / 1024,
            ('mb', 'gb'): lambda v: v / 1024,
            ('gb', 'tb'): lambda v: v / 1024,
            ('m', 'ft'): lambda v: v * 3.28084,
            ('ft', 'm'): lambda v: v * 0.3048,
            ('cm', 'in'): lambda v: v * 0.393701,
            ('in', 'cm'): lambda v: v * 2.54,
            ('hour', 'min'): lambda v: v * 60,
            ('min', 'sec'): lambda v: v * 60,
        }
        
        # Normalize unit names/aliases for better compatibility
        aliases = {
            'celsius': 'c', 'c': 'c',
            'fahrenheit': 'f', 'f': 'f',
            'kilometer': 'km', 'kilometers': 'km', 'km': 'km',
            'mile': 'mi', 'miles': 'mi', 'mi': 'mi',
            'kilogram': 'kg', 'kilograms': 'kg', 'kg': 'kg',
            'pound': 'lb', 'pounds': 'lb', 'lbs': 'lb', 'lb': 'lb',
            'meter': 'm', 'meters': 'm', 'm': 'm',
            'foot': 'ft', 'feet': 'ft', 'ft': 'ft',
            'centimeter': 'cm', 'centimeters': 'cm', 'cm': 'cm',
            'inch': 'in', 'inches': 'in', 'in': 'in',
            'hour': 'hour', 'hours': 'hour',
            'minute': 'min', 'minutes': 'min', 'min': 'min',
            'second': 'sec', 'seconds': 'sec', 'sec': 'sec',
            'kilobyte': 'kb', 'kb': 'kb',
            'megabyte': 'mb', 'mb': 'mb',
            'gigabyte': 'gb', 'gb': 'gb',
            'terabyte': 'tb', 'tb': 'tb',
        }
        
        from_norm = aliases.get(from_unit.lower(), from_unit.lower())
        to_norm = aliases.get(to_unit.lower(), to_unit.lower())
        key = (from_norm, to_norm)
        
        if key not in conversions:
            available = ', '.join(f"{a}->{b}" for a, b in conversions.keys())
            return f"Conversion {from_unit} -> {to_unit} not supported.\nAvailable: {available}"
        try:
            result = conversions[key](float(value))
            return f"{value} {from_unit} = {result:.6g} {to_unit}"
        except Exception as e:
            return f"Conversion error: {e}"

    def _detect_local_timezone(self) -> str:
        """Auto-detect the system's local IANA timezone name."""
        # Method 1: Read /etc/timezone (Debian/Ubuntu/Termux)
        try:
            with open('/etc/timezone', 'r') as f:
                tz_name = f.read().strip()
                if tz_name:
                    # Validate by trying to load it
                    from zoneinfo import ZoneInfo
                    ZoneInfo(tz_name)
                    return tz_name
        except Exception:
            pass

        # Method 2: Check TZ environment variable
        tz_env = os.environ.get('TZ', '')
        if tz_env:
            try:
                from zoneinfo import ZoneInfo
                ZoneInfo(tz_env)
                return tz_env
            except Exception:
                pass

        # Method 3: Read /etc/localtime symlink (RedHat/CentOS/Fedora)
        try:
            if os.path.exists('/etc/localtime'):
                target = os.path.realpath('/etc/localtime')
                # Usually resolves to /usr/share/zoneinfo/Asia/Jakarta
                if 'zoneinfo/' in target:
                    tz_name = target.split('zoneinfo/')[-1]
                    if tz_name:
                        from zoneinfo import ZoneInfo
                        ZoneInfo(tz_name)
                        return tz_name
        except Exception:
            pass

        # Method 4: Match UTC offset to common timezones
        try:
            local_offset = datetime.datetime.now().astimezone().utcoffset()
            offset_hours = local_offset.total_seconds() / 3600
            # Common timezone mapping
            offset_map = {
                7.0: 'Asia/Jakarta',      # WIB (Western Indonesia Time)
                8.0: 'Asia/Singapore',     # SGT / WITA
                9.0: 'Asia/Tokyo',         # JST / WIT
                0.0: 'Europe/London',       # GMT
                1.0: 'Europe/Berlin',       # CET
                5.5: 'Asia/Kolkata',        # IST
                -5.0: 'America/New_York',   # EST
                -6.0: 'America/Chicago',    # CST
                -8.0: 'America/Los_Angeles',# PST
            }
            # Find closest match (within 0.1 hours)
            for offset, name in offset_map.items():
                if abs(offset_hours - offset) < 0.1:
                    return name
            # Fallback: use Etc/GMT format
            sign = '+' if offset_hours >= 0 else ''
            return f'Etc/GMT{sign}{int(-offset_hours)}' if offset_hours != 0 else 'UTC'
        except Exception:
            pass

        return 'UTC'

    def _timestamp(self, unix_ts=None, tz='local') -> str:
        """Get current date/time with proper timezone detection synced to local terminal."""
        try:
            # Detect local timezone for 'local' parameter
            local_tz_name = self._detect_local_timezone()

            if tz == 'local' or not tz:
                tz_name = local_tz_name
            else:
                tz_name = tz

            # Parse timezone
            try:
                from zoneinfo import ZoneInfo
                tz_info = ZoneInfo(tz_name)
            except Exception:
                tz_info = None
                tz_name = local_tz_name
                try:
                    from zoneinfo import ZoneInfo
                    tz_info = ZoneInfo(tz_name)
                except Exception:
                    tz_info = None

            if unix_ts:
                dt = datetime.datetime.fromtimestamp(float(unix_ts), tz=tz_info)
            else:
                dt = datetime.datetime.now(tz=tz_info)

            # Build comprehensive output
            day_names = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
            month_names = ['January', 'February', 'March', 'April', 'May', 'June',
                          'July', 'August', 'September', 'October', 'November', 'December']

            lines = [
                f"Date: {day_names[dt.weekday()]}, {month_names[dt.month]} {dt.day}, {dt.year}",
                f"Time: {dt.strftime('%H:%M:%S')}",
                f"Timezone: {tz_name}",
                f"UTC Offset: {dt.strftime('%z') or '+0000'}",
                f"Unix Timestamp: {int(dt.timestamp())}",
            ]
            return '\n'.join(lines)
        except Exception as e:
            # Ultimate fallback: raw system time
            try:
                return f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} (local)"
            except Exception:
                return f"Error: {e}"

    def _text_transform(self, text: str, operation: str) -> str:
        ops = {
            'upper': lambda t: t.upper(),
            'lower': lambda t: t.lower(),
            'title': lambda t: t.title(),
            'strip': lambda t: t.strip(),
            'reverse': lambda t: t[::-1],
            'word_count': lambda t: f"Word count: {len(t.split())}",
            'char_count': lambda t: f"Character count: {len(t)}",
            'line_count': lambda t: f"Line count: {len(t.splitlines())}",
            'slug': lambda t: '-'.join(t.lower().split()),
            'base64_encode': lambda t: __import__('base64').b64encode(t.encode()).decode(),
            'base64_decode': lambda t: __import__('base64').b64decode(t.encode()).decode(errors='replace'),
        }
        if operation not in ops:
            return f"Unknown operation: {operation}\nAvailable: {', '.join(ops.keys())}"
        try:
            return ops[operation](text)
        except Exception as e:
            return f"Error: {e}"

    def _json_parse(self, json_string: str, query=None, pretty=True) -> str:
        try:
            data = json.loads(json_string)
            if query:
                keys = query.split('.')
                for key in keys:
                    if isinstance(data, dict):
                        data = data.get(key, f"Key '{key}' not found")
                    elif isinstance(data, list):
                        try:
                            data = data[int(key)]
                        except (ValueError, IndexError):
                            return f"Invalid index: {key}"
                    else:
                        return f"Cannot navigate into {type(data).__name__}"
            indent = 2 if pretty else None
            return json.dumps(data, indent=indent, ensure_ascii=False)
        except json.JSONDecodeError as e:
            return f"Invalid JSON: {e}"
        except Exception as e:
            return f"Error: {e}"

    def _generate_uuid(self, count: int) -> str:
        import uuid
        uuids = [str(uuid.uuid4()) for _ in range(count)]
        return '\n'.join(uuids)

    def _random_number(self, min_val, max_val, integer=True) -> str:
        try:
            min_v, max_v = float(min_val), float(max_val)
            if integer:
                return str(random.randint(int(min_v), int(max_v)))
            return str(random.uniform(min_v, max_v))
        except Exception as e:
            return f"Error: {e}"

    def _base64_tool(self, data: str, mode: str) -> str:
        import base64
        try:
            if mode == 'encode':
                return base64.b64encode(data.encode()).decode()
            elif mode == 'decode':
                return base64.b64decode(data.encode()).decode(errors='replace')
            else:
                return "Mode must be 'encode' or 'decode'"
        except Exception as e:
            return f"Error: {e}"

    def _regex_test(self, pattern: str, text: str, flags: str) -> str:
        try:
            flag_val = 0
            if 'i' in flags:
                flag_val |= re.IGNORECASE
            if 'm' in flags:
                flag_val |= re.MULTILINE
            if 's' in flags:
                flag_val |= re.DOTALL
            matches = list(re.finditer(pattern, text, flag_val))
            if not matches:
                return f"No matches for pattern '{pattern}'"
            lines = [f"Pattern '{pattern}' found {len(matches)} matches:"]
            for i, m in enumerate(matches[:10]):
                groups = [f"Group {j}: {g}" for j, g in enumerate(m.groups())]
                group_str = ' | '.join(groups) if groups else 'No groups'
                lines.append(f"  Match {i+1}: '{m.group()}' at pos {m.span()[0]}-{m.span()[1]} ({group_str})")
            if len(matches) > 10:
                lines.append(f"  ... and {len(matches) - 10} more matches")
            return '\n'.join(lines)
        except re.error as e:
            return f"Regex error: {e}"

    def _sort_data(self, text: str, mode: str, reverse: bool) -> str:
        lines = text.strip().split('\n')
        if not lines:
            return "No text to sort"
        if mode == 'numeric':
            try:
                lines.sort(key=lambda x: float(x.strip()), reverse=reverse)
            except ValueError:
                lines.sort(key=lambda x: float(''.join(c for c in x if c.isdigit() or c == '.').strip() or '0'), reverse=reverse)
        elif mode == 'length':
            lines.sort(key=len, reverse=reverse)
        elif mode == 'unique':
            seen = set()
            unique = []
            for l in lines:
                if l not in seen:
                    seen.add(l)
                    unique.append(l)
            lines = unique
        elif mode == 'reverse':
            lines = lines[::-1]
        else:
            lines.sort(key=str.lower, reverse=reverse)
        return '\n'.join(lines)

    def _count_text(self, text: str, pattern: str = None) -> str:
        lines = text.split('\n')
        words = text.split()
        chars = len(text)
        result = f"Lines: {len(lines)}\nWords: {len(words)}\nCharacters: {chars}"
        if pattern:
            count = text.lower().count(pattern.lower())
            result += f"\nPattern '{pattern}' occurrences: {count}"
        return result


    # ═══════════════════════════════════════════════════════════════
    # TOOL IMPLEMENTATIONS — PDF (v5.0)
    # ═══════════════════════════════════════════════════════════════

    def _parse_page_range(self, pages_str: str, total: int) -> list:
        """Parse page range string like '1-5', '1,3,5', 'all' into list of 0-based indices."""
        if pages_str.lower() == 'all' or not pages_str:
            return list(range(total))
        result = []
        for part in pages_str.split(','):
            part = part.strip()
            if '-' in part:
                try:
                    start, end = part.split('-', 1)
                    start = max(1, int(start))
                    end = min(total, int(end))
                    result.extend(range(start - 1, end))
                except (ValueError, IndexError):
                    pass
            else:
                try:
                    idx = int(part) - 1
                    if 0 <= idx < total:
                        result.append(idx)
                except ValueError:
                    pass
        return sorted(set(result))

    def _read_pdf(self, path: str, pages: str = 'all') -> str:
        p = Path(path).expanduser()
        if not p.exists():
            return f"Error: File not found: {path}"
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            return "Error: PyPDF2 not installed. Run: pip install PyPDF2"
        try:
            reader = PdfReader(str(p))
            total = len(reader.pages)
            page_indices = self._parse_page_range(pages, total)
            if not page_indices:
                return f"Error: No valid pages in range '{pages}' (PDF has {total} pages)"
            text = f"PDF: {p.name} | Pages: {total} | Reading: {len(page_indices)} pages\n{'='*60}\n\n"
            for idx in page_indices:
                page = reader.pages[idx]
                page_text = page.extract_text()
                if page_text:
                    text += f"--- Page {idx + 1} ---\n{page_text.strip()}\n\n"
            # Also extract metadata
            meta = reader.metadata
            if meta:
                text += f"\n{'='*60}\nMetadata:\n"
                for key in ['Title', 'Author', 'Subject', 'Creator', 'Producer', 'Pages']:
                    val = meta.get(f'/{key}', '')
                    if val:
                        text += f"  {key}: {val}\n"
            return text.strip()
        except Exception as e:
            return f"PDF read error: {e}"

    def _create_pdf(self, path: str, content: str, title: str = '') -> str:
        p = Path(path).expanduser()
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib import colors
            from reportlab.lib.units import inch
        except ImportError:
            return "Error: reportlab not installed. Run: pip install reportlab"
        try:
            p.parent.mkdir(parents=True, exist_ok=True)
            doc = SimpleDocTemplate(str(p), pagesize=letter,
                                     topMargin=0.75*inch, bottomMargin=0.75*inch,
                                     leftMargin=0.75*inch, rightMargin=0.75*inch)
            styles = getSampleStyleSheet()
            story = []
            if title:
                story.append(Paragraph(title, styles['Title']))
                story.append(Spacer(1, 12))

            lines = content.split('\n')
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                if not line:
                    story.append(Spacer(1, 6))
                elif line.startswith('#### '):
                    story.append(Paragraph(line[5:], styles['Heading4']))
                elif line.startswith('### '):
                    story.append(Paragraph(line[4:], styles['Heading3']))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], styles['Heading2']))
                elif line.startswith('# '):
                    story.append(Paragraph(line[2:], styles['Heading1']))
                elif line.startswith('- ') or line.startswith('* '):
                    story.append(Paragraph(f"  - {line[2:]}", styles['Normal']))
                elif re.match(r'^\d+\.\s', line):
                    story.append(Paragraph(f"  {line}", styles['Normal']))
                elif line.startswith('---'):
                    story.append(Spacer(1, 12))
                elif line.startswith('```'):
                    # Code block - collect all lines until closing ```
                    code_lines = []
                    i += 1
                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        code_lines.append(lines[i])
                        i += 1
                    code_text = '<br/>'.join(
                        c.replace(' ', '&nbsp;').replace('<', '&lt;').replace('>', '&gt;')
                        for c in '\n'.join(code_lines).split('\n')
                    )
                    code_style = ParagraphStyle('Code', parent=styles['Normal'],
                                                 fontName='Courier', fontSize=9,
                                                 backColor=colors.Color(0.95, 0.95, 0.95),
                                                 leftIndent=12, rightIndent=12,
                                                 spaceBefore=4, spaceAfter=4)
                    story.append(Paragraph(code_text, code_style))
                else:
                    # Support **bold** and *italic*
                    fmt = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line)
                    fmt = re.sub(r'\*(.+?)\*', r'<i>\1</i>', fmt)
                    story.append(Paragraph(fmt, styles['Normal']))
                i += 1

            doc.build(story)
            size_kb = p.stat().st_size / 1024
            return f"PDF created: {path} ({size_kb:.1f} KB)"
        except Exception as e:
            return f"PDF create error: {e}"

    def _pdf_edit(self, args: dict) -> str:
        op = args.get('operation', '')
        path = args.get('path', '')
        output = args.get('output', '')

        p = Path(path).expanduser()
        if not p.exists():
            return f"Error: File not found: {path}"

        try:
            from PyPDF2 import PdfReader, PdfWriter, PdfMerger
        except ImportError:
            return "Error: PyPDF2 not installed. Run: pip install PyPDF2"

        try:
            if op == 'info':
                reader = PdfReader(str(p))
                meta = reader.metadata
                total = len(reader.pages)
                info = f"PDF Info: {p.name}\n{'='*50}\n"
                info += f"Pages: {total}\n"
                info += f"File size: {p.stat().st_size / 1024:.1f} KB\n"
                if meta:
                    for key in ['Title', 'Author', 'Subject', 'Creator', 'Producer', 'CreationDate', 'ModDate']:
                        val = meta.get(f'/{key}', '')
                        if val:
                            info += f"{key}: {val}\n"
                # Page sizes
                if total > 0:
                    box = reader.pages[0].mediabox
                    w = float(box.width)
                    h = float(box.height)
                    info += f"Page size: {w:.0f} x {h:.0f} pts ({w/72:.1f} x {h/72:.1f} in)\n"
                return info

            elif op == 'merge':
                path2 = args.get('path2', '')
                if not path2:
                    return "Error: path2 required for merge operation"
                if not output:
                    output = str(p.parent / f"{p.stem}_merged.pdf")
                p2 = Path(path2).expanduser()
                if not p2.exists():
                    return f"Error: Second file not found: {path2}"
                merger = PdfMerger()
                merger.append(str(p))
                merger.append(str(p2))
                merger.write(output)
                merger.close()
                return f"Merged: {path} + {path2} -> {output}"

            elif op == 'split':
                if not output:
                    output = str(p.parent / f"{p.stem}_split")
                reader = PdfReader(str(p))
                out_dir = Path(output)
                out_dir.mkdir(parents=True, exist_ok=True)
                for i, page in enumerate(reader.pages):
                    writer = PdfWriter()
                    writer.add_page(page)
                    out_path = out_dir / f"page_{i+1}.pdf"
                    with open(out_path, 'wb') as f:
                        writer.write(f)
                return f"Split into {len(reader.pages)} pages: {output}/"

            elif op == 'watermark':
                text = args.get('text', 'WATERMARK')
                if not output:
                    output = str(p.parent / f"{p.stem}_watermarked.pdf")
                reader = PdfReader(str(p))
                writer = PdfWriter()
                for page in reader.pages:
                    page.merge_page(self._create_watermark(text))
                    writer.add_page(page)
                with open(output, 'wb') as f:
                    writer.write(f)
                return f"Watermarked: {output}"

            elif op == 'rotate':
                angle = args.get('angle', 90)
                if not output:
                    output = str(p.parent / f"{p.stem}_rotated.pdf")
                reader = PdfReader(str(p))
                writer = PdfWriter()
                for page in reader.pages:
                    page.rotate(int(angle))
                    writer.add_page(page)
                with open(output, 'wb') as f:
                    writer.write(f)
                return f"Rotated {angle} deg: {output}"

            elif op == 'extract':
                pages_str = args.get('pages', '')
                if not output:
                    output = str(p.parent / f"{p.stem}_extracted.pdf")
                if not pages_str:
                    return "Error: pages required for extract (e.g. '1-3' or '1,2,5')"
                reader = PdfReader(str(p))
                indices = self._parse_page_range(pages_str, len(reader.pages))
                if not indices:
                    return f"Error: No valid pages in '{pages_str}'"
                writer = PdfWriter()
                for idx in indices:
                    writer.add_page(reader.pages[idx])
                with open(output, 'wb') as f:
                    writer.write(f)
                return f"Extracted {len(indices)} pages: {output}"

            else:
                ops = 'info, merge, split, watermark, rotate, extract'
                return f"Unknown operation: {op}\nAvailable: {ops}"

        except Exception as e:
            return f"PDF edit error: {e}"

    def _create_watermark(self, text: str):
        """Create a semi-transparent watermark page."""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            buf = io.BytesIO()
            c = canvas.Canvas(buf, pagesize=letter)
            c.saveState()
            c.setFont('Helvetica', 60)
            c.setFillColorRGB(0.8, 0.8, 0.8, alpha=0.3)
            c.translate(letter[0]/2, letter[1]/2)
            c.rotate(45)
            c.drawCentredString(0, 0, text)
            c.restoreState()
            c.save()
            buf.seek(0)
            from PyPDF2 import PdfReader
            return PdfReader(buf).pages[0]
        except Exception:
            return None


    # ═══════════════════════════════════════════════════════════════
    # TOOL IMPLEMENTATIONS — DOCX (v5.0)
    # ═══════════════════════════════════════════════════════════════

    def _parse_docx_content(self, doc, content: str, tables_json: str = ''):
        """Parse markdown-like content and add to docx document."""
        lines = content.split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph('')
            elif stripped.startswith('#### '):
                doc.add_heading(stripped[5:], level=4)
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                p = doc.add_paragraph(stripped[2:], style='List Bullet')
                self._apply_inline_format(p, stripped[2:])
            elif re.match(r'^\d+\.\s', stripped):
                p = doc.add_paragraph(stripped, style='List Number')
            elif stripped.startswith('> '):
                p = doc.add_paragraph(stripped[2:], style='Quote')
            elif stripped.startswith('---'):
                doc.add_paragraph('_' * 40)
            elif stripped.startswith('```'):
                continue  # Skip code block markers in docx
            else:
                p = doc.add_paragraph(stripped)
                self._apply_inline_format(p, stripped)

        # Add tables from JSON
        if tables_json:
            try:
                tables = json.loads(tables_json)
                if isinstance(tables, list):
                    for tbl in tables:
                        headers = tbl.get('headers', [])
                        rows = tbl.get('rows', [])
                        if headers:
                            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                            table.style = 'Light Grid Accent 1'
                            for j, h in enumerate(headers):
                                table.rows[0].cells[j].text = h
                            for i, row in enumerate(rows):
                                for j, cell in enumerate(row):
                                    if j < len(headers):
                                        table.rows[i+1].cells[j].text = str(cell)
            except json.JSONDecodeError:
                pass

    def _apply_inline_format(self, paragraph, text: str):
        """Apply **bold** and *italic* formatting to a paragraph."""
        # Clear existing runs and re-add with formatting
        p = paragraph
        p.clear()
        # Simple regex-based inline formatting
        parts = re.split(r'(\*\*\w.*?\*\*|\*\w.*?\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                p.add_run(part[1:-1]).italic = True
            else:
                p.add_run(part)

    def _read_docx(self, path: str) -> str:
        p = Path(path).expanduser()
        if not p.exists():
            return f"Error: File not found: {path}"
        try:
            from docx import Document
        except ImportError:
            return "Error: python-docx not installed. Run: pip install python-docx"
        try:
            doc = Document(str(p))
            result = []
            result.append(f"DOCX: {p.name} | Size: {p.stat().st_size / 1024:.1f} KB")
            result.append("=" * 60)

            # Core properties
            try:
                cp = doc.core_properties
                props = []
                if cp.title: props.append(f"Title: {cp.title}")
                if cp.author: props.append(f"Author: {cp.author}")
                if cp.subject: props.append(f"Subject: {cp.subject}")
                if cp.created: props.append(f"Created: {cp.created}")
                if cp.modified: props.append(f"Modified: {cp.modified}")
                if props:
                    result.append("Properties: " + " | ".join(props))
            except Exception:
                pass

            result.append("")

            # Paragraphs
            for para in doc.paragraphs:
                style_name = para.style.name if para.style else ''
                text = para.text
                if not text:
                    continue
                if 'Heading' in style_name:
                    level = style_name.replace('Heading ', '').replace('Heading', '1')
                    try:
                        level_int = int(level)
                        result.append(f"\n{'#' * level_int} {text}")
                    except ValueError:
                        result.append(f"\n# {text}")
                elif 'Bullet' in style_name or 'List' in style_name:
                    result.append(f"  - {text}")
                else:
                    result.append(text)

            # Tables
            for ti, table in enumerate(doc.tables):
                result.append(f"\n[Table {ti+1}]")
                for row in table.rows:
                    cells = [cell.text.replace('\n', ' | ') for cell in row.cells]
                    result.append('  | ' + ' | '.join(cells) + ' |')

            return '\n'.join(result)
        except Exception as e:
            return f"DOCX read error: {e}"

    def _create_docx(self, path: str, content: str, title: str = '', tables: str = '') -> str:
        p = Path(path).expanduser()
        try:
            from docx import Document
        except ImportError:
            return "Error: python-docx not installed. Run: pip install python-docx"
        try:
            p.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
            if title:
                doc.add_heading(title, 0)
            self._parse_docx_content(doc, content, tables)
            doc.save(str(p))
            size_kb = p.stat().st_size / 1024
            return f"DOCX created: {path} ({size_kb:.1f} KB)"
        except Exception as e:
            return f"DOCX create error: {e}"

    def _docx_info(self, path: str) -> str:
        p = Path(path).expanduser()
        if not p.exists():
            return f"Error: File not found: {path}"
        try:
            from docx import Document
        except ImportError:
            return "Error: python-docx not installed. Run: pip install python-docx"
        try:
            doc = Document(str(p))
            info = []
            info.append(f"DOCX Info: {p.name}")
            info.append("=" * 50)
            info.append(f"File size: {p.stat().st_size / 1024:.1f} KB")

            # Core properties
            try:
                cp = doc.core_properties
                info.append(f"Title: {cp.title or '(none)'}")
                info.append(f"Author: {cp.author or '(none)'}")
                info.append(f"Subject: {cp.subject or '(none)'}")
                info.append(f"Category: {cp.category or '(none)'}")
                info.append(f"Keywords: {cp.keywords or '(none)'}")
                info.append(f"Created: {cp.created}")
                info.append(f"Modified: {cp.modified}")
                info.append(f"Last printed: {cp.last_printed_by or '(none)'}")
                info.append(f"Revision: {cp.revision}")
            except Exception as e:
                info.append(f"Properties error: {e}")

            # Content stats
            para_count = len(doc.paragraphs)
            word_count = sum(len(p.text.split()) for p in doc.paragraphs)
            char_count = sum(len(p.text) for p in doc.paragraphs)
            table_count = len(doc.tables)
            heading_count = sum(1 for p in doc.paragraphs if 'Heading' in (p.style.name if p.style else ''))

            info.append("")
            info.append("Content Statistics:")
            info.append(f"  Paragraphs: {para_count}")
            info.append(f"  Words: {word_count}")
            info.append(f"  Characters: {char_count}")
            info.append(f"  Headings: {heading_count}")
            info.append(f"  Tables: {table_count}")

            # Sections info
            sections = doc.sections
            if sections:
                sec = sections[0]
                info.append("")
                info.append("Page Setup:")
                info.append(f"  Width: {sec.page_width.inches:.1f} in")
                info.append(f"  Height: {sec.page_height.inches:.1f} in")
                info.append(f"  Margins: L={sec.left_margin.inches:.1f} R={sec.right_margin.inches:.1f} T={sec.top_margin.inches:.1f} B={sec.bottom_margin.inches:.1f}")

            return '\n'.join(info)
        except Exception as e:
            return f"DOCX info error: {e}"

    def _todowrite(self, title: str, items: list, path: str = '') -> str:
        title = title.lstrip('#').strip()
        lines = [f'# {title}']
        lines.append('')
        for i, item in enumerate(items):
            text = item.get('text', '')
            if isinstance(text, str):
                text = text.strip().strip("'").strip('"')
            done = item.get('done', False)
            if done:
                lines.append(f'[✓] {text}')
            else:
                lines.append(f'[ ] {text}')
        # Save to memory
        if self._memory:
            self._memory.set_todo(items)
        if path:
            try:
                p = Path(path).expanduser()
                p.parent.mkdir(parents=True, exist_ok=True)
                p.write_text('\n'.join(lines), encoding='utf-8')
                lines.append('')
                lines.append(f'Saved to: {path}')
            except Exception as e:
                lines.append(f'Save error: {e}')
        return '\n'.join(lines)

    def _todolist_get(self) -> str:
        if not self._memory or not self._memory.todo_items:
            return '# No todo list set. Use todowrite to create one first.'
        lines = ['# Active Todo List']
        for i, item in enumerate(self._memory.todo_items):
            text = item.get('text', '')
            done = item.get('done', False)
            status = '✓' if done else ' '
            lines.append(f'  [{status}] [{i}] {text}')
        return '\n'.join(lines) + '\n\nUse todolist_update with an index to mark items as done.'

    def _todolist_update(self, index: int, done: bool = None, text: str = '') -> str:
        if not self._memory or not self._memory.todo_items:
            return 'Error: No todo list in memory.'
        if index < 0 or index >= len(self._memory.todo_items):
            return f'Error: Index {index} out of range (0-{len(self._memory.todo_items) - 1})'
        item = self._memory.todo_items[index]
        if done is not None:
            item['done'] = done
        if text:
            item['text'] = text
        status = '✓' if item['done'] else '○'
        return f'{status} {item["text"]}'

    def _todolist_update_all(self) -> str:
        if not self._memory or not self._memory.todo_items:
            return 'Error: No todo list in memory.'
        count = 0
        for item in self._memory.todo_items:
            if not item.get('done'):
                item['done'] = True
                count += 1
        total = len(self._memory.todo_items)
        if count == 0:
            return f'✓ All {total} tasks already completed!'
        return f'✓ Marked {count}/{total} item(s) as done. All {total} tasks completed!'


    # ═══════════════════════════════════════════════════════════════
    # TOOL IMPLEMENTATIONS — IMAGE (v5.0)
    # ═══════════════════════════════════════════════════════════════

    def _image_view(self, path: str, width: int = 80) -> str:
        p = Path(path).expanduser()
        if not p.exists():
            return f"Error: File not found: {path}"
        try:
            from PIL import Image
        except ImportError:
            return "Error: Pillow not installed. Run: pip install Pillow"
        try:
            img = Image.open(str(p))
            orig_w, orig_h = img.size
            fmt = img.format or 'Unknown'
            mode = img.mode

            # Resize for ASCII (maintain aspect ratio, chars are ~2x tall)
            ratio = (orig_h / orig_w) * 0.5
            new_w = min(width, orig_w)
            new_h = max(1, int(new_w * ratio))
            new_h = min(new_h, 60)  # Max terminal height

            img_resized = img.resize((new_w, new_h), Image.LANCZOS if hasattr(Image, 'LANCZOS') else Image.BILINEAR)
            img_gray = img_resized.convert('L')

            # ASCII chars (dark to light)
            ASCII_CHARS = ' .\'`^",:;Il!i><~+_-?][}{1)(|\\/tfjrxnuvczXYUJCLQ0OZmwqpdbkhao*#MW&8%B@$'

            ascii_lines = []
            for y in range(new_h):
                row = ''
                for x in range(new_w):
                    pixel = img_gray.getpixel((x, y))
                    idx = pixel * (len(ASCII_CHARS) - 1) // 255
                    row += ASCII_CHARS[idx]
                ascii_lines.append(row)

            header = f"Image: {p.name} | {orig_w}x{orig_h} | {fmt} | {mode}\n"
            header += f"Size: {p.stat().st_size / 1024:.1f} KB\n"
            header += "=" * width + "\n"

            return header + '\n'.join(ascii_lines)
        except Exception as e:
            return f"Image view error: {e}"

    def _image_info(self, path: str) -> str:
        p = Path(path).expanduser()
        if not p.exists():
            return f"Error: File not found: {path}"
        try:
            from PIL import Image
        except ImportError:
            return "Error: Pillow not installed. Run: pip install Pillow"
        try:
            img = Image.open(str(p))
            w, h = img.size
            info = []
            info.append(f"Image Info: {p.name}")
            info.append("=" * 50)
            info.append(f"Format: {img.format or 'Unknown'}")
            info.append(f"Mode: {img.mode} (colorspace)")
            info.append(f"Size: {w} x {h} pixels")
            info.append(f"Megapixels: {w * h / 1e6:.2f} MP")
            info.append(f"Aspect ratio: {w/h:.3f}" if h else "Aspect ratio: N/A")
            info.append(f"File size: {p.stat().st_size / 1024:.1f} KB")

            if img.mode == 'RGBA' and 'transparency' in img.info:
                info.append("Transparency: Yes")

            # DPI
            dpi = img.info.get('dpi')
            if dpi:
                info.append(f"DPI: {dpi[0]} x {dpi[1]}")

            # Animation (GIF)
            if getattr(img, 'n_frames', 1) > 1:
                info.append(f"Frames: {img.n_frames}")
                try:
                    info.append(f"Duration: {img.info.get('duration', '?')} ms per frame")
                except Exception:
                    pass

            # EXIF
            try:
                exif = img._getexif()
                if exif:
                    info.append("")
                    info.append("EXIF Data:")
                    from PIL.ExifTags import TAGS
                    for tag_id, value in exif.items():
                        tag = TAGS.get(tag_id, tag_id)
                        if isinstance(value, bytes):
                            continue
                        info.append(f"  {tag}: {value}")
            except Exception:
                pass

            # Color statistics
            try:
                img_rgb = img.convert('RGB')
                pixels = list(img_rgb.getdata())
                r_avg = sum(px[0] for px in pixels) / len(pixels)
                g_avg = sum(px[1] for px in pixels) / len(pixels)
                b_avg = sum(px[2] for px in pixels) / len(pixels)
                info.append("")
                info.append("Color Averages:")
                info.append(f"  Red: {r_avg:.0f} | Green: {g_avg:.0f} | Blue: {b_avg:.0f}")
                brightness = (r_avg * 299 + g_avg * 587 + b_avg * 114) / 1000
                info.append(f"  Brightness: {brightness:.0f}/255")
            except Exception:
                pass

            return '\n'.join(info)
        except Exception as e:
            return f"Image info error: {e}"


    # ═══════════════════════════════════════════════════════════════
    # TOOL IMPLEMENTATIONS — OCR (v6.1)
    # pytesseract (primary) + easyocr (fallback)
    # ═══════════════════════════════════════════════════════════════

    @staticmethod
    def _ocr_preprocess_image(img, mode: str):
        """
        Apply preprocessing to a PIL Image to improve OCR accuracy.
        Modes: none, grayscale, threshold, denoise, sharpen, enhance.
        Returns a processed PIL Image.
        """
        from PIL import Image, ImageFilter, ImageOps

        if mode == 'none' or not mode:
            return img

        # Ensure RGB mode for processing
        if img.mode == 'RGBA':
            img = img.convert('RGB')
        elif img.mode != 'RGB':
            img = img.convert('RGB')

        if mode == 'grayscale':
            return img.convert('L')

        elif mode == 'threshold':
            gray = img.convert('L')
            # Otsu-like simple threshold
            threshold = 128
            return gray.point(lambda x: 255 if x > threshold else 0, '1')

        elif mode == 'denoise':
            return img.filter(ImageFilter.MedianFilter(size=3))

        elif mode == 'sharpen':
            return img.filter(ImageFilter.SHARPEN)

        elif mode == 'enhance':
            from PIL import ImageEnhance
            img = ImageEnhance.Contrast(img).enhance(1.5)
            img = ImageEnhance.Sharpness(img).enhance(1.5)
            img = ImageEnhance.Brightness(img).enhance(1.1)
            return img

        return img

    def _ocr_with_tesseract(self, img, language: str) -> str:
        """Run OCR using pytesseract. Returns extracted text."""
        try:
            import pytesseract
        except ImportError:
            return None  # Signal: pytesseract not available

        try:
            # pytesseract expects 'eng', 'eng+ind', etc.
            text = pytesseract.image_to_string(img, lang=language)
            return text.strip() if text.strip() else None
        except pytesseract.TesseractNotFoundError:
            return None  # Tesseract binary not found
        except Exception:
            return None  # Other tesseract error

    def _ocr_with_easyocr(self, img, language: str) -> str:
        """Run OCR using EasyOCR. Returns extracted text."""
        try:
            # pyrefly: ignore [missing-import]
            import easyocr
        except ImportError:
            return None  # Signal: easyocr not available

        try:
            # EasyOCR uses different language codes
            # Map common tesseract codes to easyocr codes
            lang_map = {
                'eng': 'en', 'ind': 'id', 'jpn': 'ja', 'kor': 'ko',
                'chi_sim': 'ch_sim', 'chi_tra': 'ch_tra', 'spa': 'es',
                'fra': 'fr', 'deu': 'de', 'por': 'pt', 'rus': 'ru',
                'ara': 'ar', 'hin': 'hi', 'tha': 'th', 'vie': 'vi',
            }
            # Handle compound languages like 'eng+ind'
            langs = language.split('+')
            easy_langs = []
            for lang in langs:
                easy_langs.append(lang_map.get(lang.strip(), lang.strip()))

            reader = easyocr.Reader(easy_langs, gpu=False, verbose=False)

            # Convert PIL Image to numpy array
            import numpy as np
            img_array = np.array(img)

            results = reader.readtext(img_array)
            if not results:
                return None

            # Concatenate all detected text
            lines = []
            for detection in results:
                text = detection[1]
                confidence = detection[2]
                if confidence > 0.3:  # Filter low-confidence results
                    lines.append(text)

            full_text = '\n'.join(lines) if lines else None
            return full_text
        except Exception:
            return None

    def _ocr_read(self, path: str, language: str = 'eng',
                  engine: str = 'auto', preprocess: str = 'none') -> str:
        """
        Extract text from an image file using OCR.
        Supports pytesseract and easyocr engines with automatic fallback.
        """
        p = Path(path).expanduser()
        if not p.exists():
            return f"Error: File not found: {path}"

        # Check if Pillow is available
        try:
            from PIL import Image
        except ImportError:
            return "Error: Pillow not installed. Run: pip install Pillow"

        try:
            img = Image.open(str(p))
        except Exception as e:
            return f"Error: Cannot open image: {e}"

        # Apply preprocessing
        img = self._ocr_preprocess_image(img, preprocess)

        # Determine engine strategy
        engine = engine.lower().strip() if engine else 'auto'

        if engine == 'tesseract':
            result = self._ocr_with_tesseract(img, language)
            if result:
                return f"OCR Result (Tesseract | lang: {language} | preprocess: {preprocess}):\n{'=' * 60}\n{result}"
            return ("Error: Tesseract failed. Make sure pytesseract and tesseract-ocr are installed.\n"
                    "  pip install pytesseract\n"
                    "  Termux: pkg install tesseract\n"
                    "  Linux:  sudo apt install tesseract-ocr\n"
                    "  macOS:  brew install tesseract\n"
                    "  Or try engine=\"easyocr\" (pip install easyocr)")

        elif engine == 'easyocr':
            result = self._ocr_with_easyocr(img, language)
            if result:
                return f"OCR Result (EasyOCR | lang: {language} | preprocess: {preprocess}):\n{'=' * 60}\n{result}"
            return ("Error: EasyOCR failed. Make sure easyocr is installed.\n"
                    "  pip install easyocr\n"
                    "  Note: First run downloads language models (~100MB).\n"
                    "  Or try engine=\"tesseract\" (pip install pytesseract)")

        else:  # engine == 'auto'
            # Try tesseract first, then easyocr
            tesseract_result = self._ocr_with_tesseract(img, language)
            if tesseract_result:
                return f"OCR Result (Tesseract | lang: {language} | preprocess: {preprocess}):\n{'=' * 60}\n{tesseract_result}"

            # Fallback to easyocr
            easyocr_result = self._ocr_with_easyocr(img, language)
            if easyocr_result:
                return f"OCR Result (EasyOCR | lang: {language} | preprocess: {preprocess}):\n{'=' * 60}\n{easyocr_result}"

            # Both failed
            return ("Error: No OCR engine available. Install one of:\n\n"
                    "  Option 1 - Tesseract (recommended, fast):\n"
                    "    pip install pytesseract\n"
                    "    Termux: pkg install tesseract\n"
                    "    Linux:  sudo apt install tesseract-ocr\n"
                    "    macOS:  brew install tesseract\n\n"
                    "  Option 2 - EasyOCR (no system dependency, slower):\n"
                    "    pip install easyocr\n"
                    "    Note: Downloads ~100MB language models on first use.")

    def _ocr_url(self, url: str, language: str = 'eng',
                 engine: str = 'auto', preprocess: str = 'none') -> str:
        """
        Download an image from URL and extract text using OCR.
        Downloads to a temporary file, then calls _ocr_read.
        """
        try:
            import httpx
        except ImportError:
            return "Error: httpx not installed. Run: pip install httpx"

        try:
            with httpx.Client(timeout=30, follow_redirects=True) as client:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Linux; Android 14) AppleWebKit/537.36 Chrome/125.0.0.0 Safari/537.36'
                }
                response = client.get(url, headers=headers)
                response.raise_for_status()

            # Determine file extension from URL or content-type
            content_type = response.headers.get('content-type', '')
            ext = '.jpg'
            if 'png' in content_type:
                ext = '.png'
            elif 'webp' in content_type:
                ext = '.webp'
            elif 'gif' in content_type:
                ext = '.gif'
            elif 'bmp' in content_type:
                ext = '.bmp'
            else:
                # Try to guess from URL
                from urllib.parse import urlparse
                parsed = urlparse(url)
                path_lower = parsed.path.lower()
                for e in ['.png', '.jpg', '.jpeg', '.webp', '.gif', '.bmp', '.tiff']:
                    if path_lower.endswith(e):
                        ext = e
                        break

            # Save to temp file
            tmp = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
            tmp.write(response.content)
            tmp.close()

            try:
                # Run OCR
                result = self._ocr_read(tmp.name, language, engine, preprocess)
                if not result.startswith('Error'):
                    # Add source info
                    result = f"Source: {url}\nSize: {len(response.content):,} bytes\n{result}"
                return result
            finally:
                # Clean up temp file
                try:
                    os.unlink(tmp.name)
                except Exception:
                    pass

        except httpx.TimeoutException:
            return f"Error: Download timed out for {url}"
        except httpx.HTTPStatusError as e:
            return f"Error: HTTP {e.response.status_code} when downloading {url}"
        except Exception as e:
            return f"Error downloading image: {e}"

    # ═══════════════════════════════════════════════════════════════
    # TOOL IMPLEMENTATIONS — VIDEO (v5.0)
    # ═══════════════════════════════════════════════════════════════

    def _video_info(self, path: str) -> str:
        p = Path(path).expanduser()
        if not p.exists():
            return f"Error: File not found: {path}"
        # Check ffprobe
        ffprobe = shutil.which('ffprobe')
        if not ffprobe:
            return ("Error: ffprobe not found. Install ffmpeg:\n"
                    "  Termux: pkg install ffmpeg\n"
                    "  Linux:  sudo apt install ffmpeg\n"
                    "  macOS:  brew install ffmpeg")
        try:
            result = subprocess.run(
                [ffprobe, '-v', 'quiet', '-print_format', 'json',
                 '-show_format', '-show_streams', str(p)],
                capture_output=True, text=True, timeout=15
            )
            if result.returncode != 0:
                return f"ffprobe error: {result.stderr[:500]}"

            data = json.loads(result.stdout)
            info = []
            info.append(f"Video Info: {p.name}")
            info.append("=" * 50)
            info.append(f"File size: {p.stat().st_size / (1024*1024):.2f} MB")

            fmt = data.get('format', {})
            if fmt:
                duration = float(fmt.get('duration', 0))
                mins, secs = divmod(int(duration), 60)
                hours, mins = divmod(mins, 60)
                dur_str = f"{hours}:{mins:02d}:{secs:02d}" if hours else f"{mins}:{secs:02d}"
                info.append(f"Duration: {dur_str}")
                info.append(f"Bitrate: {int(fmt.get('bit_rate', 0)) / 1000:.0f} kbps")
                nb_streams = fmt.get('nb_streams', '?')
                info.append(f"Streams: {nb_streams}")
                if fmt.get('tags'):
                    for k, v in fmt['tags'].items():
                        info.append(f"{k}: {v}")

            streams = data.get('streams', [])
            for stream in streams:
                stype = stream.get('codec_type', 'unknown')
                info.append("")
                info.append(f"[{stype.upper()} Stream]")
                if stype == 'video':
                    codec = stream.get('codec_name', '?')
                    w = stream.get('width', '?')
                    h = stream.get('height', '?')
                    fps_raw = stream.get('r_frame_rate', '0/1')
                    try:
                        num, den = fps_raw.split('/')
                        fps = float(num) / float(den) if float(den) else 0
                    except (ValueError, ZeroDivisionError):
                        fps = 0
                    info.append(f"  Codec: {codec}")
                    info.append(f"  Resolution: {w}x{h}")
                    info.append(f"  FPS: {fps:.2f}")
                    if stream.get('bit_rate'):
                        info.append(f"  Bitrate: {int(stream['bit_rate']) / 1000:.0f} kbps")
                    if stream.get('pix_fmt'):
                        info.append(f"  Pixel format: {stream['pix_fmt']}")
                    sar = stream.get('sample_aspect_ratio', '')
                    dar = stream.get('display_aspect_ratio', '')
                    if sar and sar != 'N/A':
                        info.append(f"  SAR: {sar}  DAR: {dar}")
                    if stream.get('tags'):
                        for k, v in stream['tags'].items():
                            info.append(f"  Tag: {k} = {v}")
                elif stype == 'audio':
                    codec = stream.get('codec_name', '?')
                    sample_rate = stream.get('sample_rate', '?')
                    channels = stream.get('channels', '?')
                    info.append(f"  Codec: {codec}")
                    info.append(f"  Sample rate: {sample_rate} Hz")
                    info.append(f"  Channels: {channels}")
                    if stream.get('bit_rate'):
                        info.append(f"  Bitrate: {int(stream['bit_rate']) / 1000:.0f} kbps")
                    if stream.get('tags'):
                        for k, v in stream['tags'].items():
                            if k.lower() in ('title', 'language', 'handler_name'):
                                info.append(f"  {k}: {v}")

            return '\n'.join(info)
        except subprocess.TimeoutExpired:
            return "Error: ffprobe timed out (file too large?)"
        except json.JSONDecodeError:
            return f"Error: Could not parse ffprobe output"
        except Exception as e:
            return f"Video info error: {e}"

    def _video_play(self, path: str) -> str:
        p = Path(path).expanduser()
        if not p.exists():
            return f"Error: File not found: {path}"
        # Try different open commands
        openers = []
        if shutil.which('termux-open'):
            openers.append(('termux-open', ['termux-open', str(p)]))
        if shutil.which('xdg-open'):
            openers.append(('xdg-open', ['xdg-open', str(p)]))
        if shutil.which('open'):
            openers.append(('open', ['open', str(p)]))
        if shutil.which('am'):  # Android activity manager
            openers.append(('am', ['am', 'start', '-a', 'android.intent.action.VIEW',
                                    '-d', f'file://{p.resolve()}', '-t', 'video/*']))

        if not openers:
            return ("Error: No media player found.\n"
                    "Install one of: termux-api (Termux), xdg-utils (Linux)\n"
                    "  Termux: pkg install termux-api\n"
                    "  Linux:  sudo apt install xdg-utils")
        name, cmd = openers[0]
        try:
            subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            return f"Opened with {name}: {p.name}"
        except Exception as e:
            return f"Error opening: {e}"


    # ═══════════════════════════════════════════════════════════════
    # TOOL IMPLEMENTATIONS — APK (v5.0 BETA)
    # ═══════════════════════════════════════════════════════════════

    def _extract_axml_strings(self, data: bytes) -> list:
        """Extract readable strings from Android Binary XML (AndroidManifest.xml)."""
        strings = []
        # Scan for UTF-16LE string patterns in the binary data
        i = 0
        current = ""
        while i < len(data) - 1:
            b1 = data[i]
            b2 = data[i + 1]
            char_code = b1 | (b2 << 8)

            if 32 <= char_code <= 126:
                current += chr(char_code)
            elif char_code == 0 or char_code > 126:
                if len(current) > 2:
                    if any(kw in current for kw in [
                        '.', 'permission', 'activity', 'service', 'receiver',
                        'provider', 'application', 'manifest', 'uses-',
                        'intent', 'action', 'category', 'version', 'package',
                        'android', 'com.', 'org.', 'net.', 'io.', 'gov.',
                        'minSdk', 'targetSdk', 'debuggable', 'exported',
                        'enabled', 'name', 'value', 'scheme', 'host', 'path',
                        'portrait', 'landscape', 'configChanges', 'theme',
                        'label', 'icon', 'backup', 'supportsRtl'
                    ]):
                        strings.append(current)
                current = ""
            else:
                current = ""
            i += 2
        return strings

    def _apk_analyze(self, path: str) -> str:
        p = Path(path).expanduser()
        if not p.exists():
            return f"Error: File not found: {path}"
        if not zipfile.is_zipfile(str(p)):
            return "Error: Not a valid APK file (APK files are ZIP archives)"

        info = []
        info.append(f"APK Analysis [BETA]: {p.name}")
        info.append("=" * 60)
        info.append(f"File size: {p.stat().st_size / (1024*1024):.2f} MB")

        try:
            with zipfile.ZipFile(str(p), 'r') as zf:
                all_files = zf.namelist()
                info.append(f"Total files: {len(all_files)}")

                # Directory structure
                dirs = {}
                for f in all_files:
                    parts = f.split('/')
                    if len(parts) > 1:
                        d = parts[0] + '/'
                        dirs[d] = dirs.get(d, 0) + 1
                info.append("")
                info.append("Directory Structure:")
                for d, count in sorted(dirs.items()):
                    info.append(f"  {d} ({count} files)")

                # Parse AndroidManifest.xml
                info.append("")
                info.append("AndroidManifest.xml (extracted strings):")
                try:
                    manifest_data = zf.read('AndroidManifest.xml')
                    strings = self._extract_axml_strings(manifest_data)

                    # Categorize strings
                    packages = []
                    permissions = []
                    activities = []
                    services = []
                    receivers = []
                    providers = []
                    other = []

                    for s in strings:
                        sl = s.lower()
                        if 'uses-permission' in sl or 'permission' == sl.split('.')[-1]:
                            permissions.append(s)
                        elif 'activity' in sl:
                            activities.append(s)
                        elif 'service' in sl:
                            services.append(s)
                        elif 'receiver' in sl:
                            receivers.append(s)
                        elif 'provider' in sl:
                            providers.append(s)
                        elif re.match(r'^[a-z][a-z0-9_]*(\.[a-z][a-z0-9_]*)+', sl):
                            packages.append(s)
                        else:
                            other.append(s)

                    if packages:
                        info.append("  Package/App:")
                        for s in packages:
                            info.append(f"    {s}")
                    if permissions:
                        info.append(f"  Permissions ({len(permissions)}):")
                        for s in permissions:
                            info.append(f"    {s}")
                    if activities:
                        info.append(f"  Activities ({len(activities)}):")
                        for s in activities:
                            info.append(f"    {s}")
                    if services:
                        info.append(f"  Services ({len(services)}):")
                        for s in services:
                            info.append(f"    {s}")
                    if receivers:
                        info.append(f"  Receivers ({len(receivers)}):")
                        for s in receivers:
                            info.append(f"    {s}")
                    if providers:
                        info.append(f"  Providers ({len(providers)}):")
                        for s in providers:
                            info.append(f"    {s}")

                except KeyError:
                    info.append("  (AndroidManifest.xml not found)")

                # DEX files
                info.append("")
                info.append("DEX Files:")
                dex_count = 0
                for f in all_files:
                    if f.endswith('.dex'):
                        fi = zf.getinfo(f)
                        info.append(f"  {f} ({fi.file_size / 1024:.1f} KB)")
                        dex_count += 1
                if not dex_count:
                    info.append("  (No DEX files found)")

                # Native libraries
                info.append("")
                info.append("Native Libraries (.so):")
                so_count = 0
                so_archs = set()
                for f in all_files:
                    if f.endswith('.so'):
                        so_count += 1
                        parts = f.split('/')
                        if 'lib/' in f:
                            idx = parts.index('lib')
                            if idx + 1 < len(parts):
                                so_archs.add(parts[idx + 1])
                if so_count:
                    info.append(f"  Total: {so_count} native libraries")
                    if so_archs:
                        info.append(f"  Architectures: {', '.join(sorted(so_archs))}")
                    # List some .so files
                    for f in all_files:
                        if f.endswith('.so'):
                            fi = zf.getinfo(f)
                            info.append(f"  {f} ({fi.file_size / 1024:.1f} KB)")
                else:
                    info.append("  (No native libraries)")

                # Certificates (META-INF)
                info.append("")
                info.append("Signing (META-INF):")
                cert_files = [f for f in all_files if f.startswith('META-INF/')]
                if cert_files:
                    for f in cert_files:
                        fi = zf.getinfo(f)
                        info.append(f"  {f} ({fi.file_size / 1024:.1f} KB)")
                else:
                    info.append("  (Not signed or signature stripped)")

                # resources.arsc
                info.append("")
                arsc = [f for f in all_files if 'resources.arsc' in f]
                if arsc:
                    fi = zf.getinfo(arsc[0])
                    info.append(f"Resources: {arsc[0]} ({fi.file_size / 1024:.1f} KB)")

                # Top 10 largest files
                info.append("")
                info.append("Top 10 Largest Files:")
                sized = [(f, zf.getinfo(f).file_size) for f in all_files if not f.endswith('/')]
                sized.sort(key=lambda x: x[1], reverse=True)
                for f, sz in sized[:10]:
                    info.append(f"  {f} ({sz / 1024:.1f} KB)")

        except Exception as e:
            info.append(f"Analysis error: {e}")

        info.append("")
        info.append("[BETA] Note: Binary XML parsing has limitations.")
        info.append("For full analysis, use jadx, apktool, or androguard.")

        return '\n'.join(info)

    # ═══════════════════════════════════════════════════════════════
    # LIVE SEARCH & MODEL SEARCH IMPLEMENTATIONS (v5.2)
    # ═══════════════════════════════════════════════════════════════

    def _live_search(self, query: str, max_results: int, source: str) -> str:
        """Live web search with multiple sources and fallback."""
        all_results = []
        sources_used = []

        # Source 1: DuckDuckGo
        if source in ('all', 'duckduckgo'):
            try:
                from ddgs import DDGS
                ddgs_results = []
                with DDGS() as ddgs:
                    for r in ddgs.text(query, max_results=max_results):
                        ddgs_results.append({
                            'title': r.get('title', ''),
                            'url': r.get('href', ''),
                            'body': r.get('body', ''),
                            'source': 'DuckDuckGo'
                        })
                all_results.extend(ddgs_results)
                if ddgs_results:
                    sources_used.append(f'DuckDuckGo ({len(ddgs_results)})')
            except Exception as e:
                all_results.append({
                    'title': f'[DuckDuckGo Error]',
                    'url': '',
                    'body': str(e),
                    'source': 'DuckDuckGo'
                })

        # Source 2: Google News (via DuckDuckGo news)
        if source in ('all', 'google'):
            try:
                try:
                    from ddgs import DDGS
                except ImportError:
                    from duckduckgo_search import DDGS
                news_results = []
                with DDGS() as ddgs:
                    for r in ddgs.news(query, max_results=max(max_results, 3)):
                        news_results.append({
                            'title': r.get('title', ''),
                            'url': r.get('url', '') or r.get('href', ''),
                            'body': r.get('body', ''),
                            'source': 'Google News'
                        })
                all_results.extend(news_results)
                if news_results:
                    sources_used.append(f'Google News ({len(news_results)})')
            except Exception:
                pass

        # Source 3: Bing (via DuckDuckGo with different region)
        if source in ('all', 'bing') and len(all_results) < 3:
            try:
                import httpx
                # Use Bing search via URL construction
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
                }
                r = httpx.get(f'https://www.bing.com/search?q={query.replace(" ", "+")}&count={max_results}',
                             headers=headers, timeout=15, follow_redirects=True)
                if r.status_code == 200:
                    import re as _re
                    # Extract titles and snippets from Bing HTML
                    titles = _re.findall(r'<a[^>]+class="[^"]*"[^>]*>([^<]+)</a>', r.text[:20000])
                    snippets = _re.findall(r'<span[^>]*>([^<]{50,300})</span>', r.text[:20000])
                    bing_count = min(len(titles), max_results)
                    for i in range(bing_count):
                        title = titles[i].strip() if i < len(titles) else ''
                        body = snippets[i].strip() if i < len(snippets) else ''
                        if title and len(title) > 5:
                            all_results.append({
                                'title': title,
                                'url': f'https://www.bing.com/search?q={query.replace(" ", "+")}',
                                'body': body[:200],
                                'source': 'Bing'
                            })
                    if bing_count > 0:
                        sources_used.append(f'Bing ({bing_count})')
            except Exception:
                pass

        if not all_results:
            return f"No results found for: {query}\n\nTry: /live_search {query}"

        # Format results
        header = f"Live Search: '{query}'"
        if sources_used:
            header += f"  |  Sources: {', '.join(sources_used)}"
        header += f"  |  Total: {len(all_results)} results"

        lines = [header, '=' * min(len(header), 70)]
        for i, r in enumerate(all_results[:max_results * 2], 1):
            title = r.get('title', '').strip()
            url = r.get('url', '').strip()
            body = r.get('body', '').strip()
            src = r.get('source', '').strip()
            if not title:
                continue
            line = f"\n{i}. {title}"
            if url:
                line += f"\n   {url}"
            if body:
                line += f"\n   {body[:200]}"
            if src:
                line += f"\n   [Source: {src}]"
            lines.append(line)

        return '\n'.join(lines)

    def _search_models(self, query: str, max_results: int) -> str:
        """Search models from the current provider's API."""
        try:
            # Import config and providers
            from .config import cfg
            from .providers import create_provider

            pid = cfg.active_provider
            pconfig = cfg.get_provider_config(pid)
            api_key = cfg.get_api_key(pid)

            if not api_key:
                return "Error: No API key set for current provider. Use /key to set one."

            provider = create_provider(pid, pconfig, api_key)
            models = provider.fetch_models()

            if not models:
                # Fallback to popular models from config
                popular = pconfig.get('popular_models', [])
                if query:
                    filtered = [m for m in popular if query.lower() in m.lower()]
                else:
                    filtered = popular
                if filtered:
                    result = f"Models from config ({len(filtered)} found):\n\n"
                    for i, m in enumerate(filtered[:max_results], 1):
                        result += f"  {i}. {m}\n"
                    result += f"\n  [Source: config (API fetch returned empty)]"
                    return result
                return f"No models found for '{query}'. Try /live_models to see all available models."

            # Filter by query
            if query:
                q = query.lower()
                filtered = [m for m in models
                           if q in m.get('id', '').lower()
                           or q in m.get('name', '').lower()]
            else:
                filtered = models

            if not filtered:
                return f"No models matching '{query}'. Available: {len(models)} total models."

            # Sort: free first, then alphabetically
            filtered.sort(key=lambda x: (not x.get('free', False), x.get('id', '').lower()))
            filtered = filtered[:max_results]

            provider_name = pconfig.get('name', pid)
            result = f"Live Models from {provider_name} ({len(filtered)} shown"
            if query:
                result += f", filtered by '{query}'"
            result += f"):\n\n"

            for i, m in enumerate(filtered, 1):
                mid = m.get('id', '')
                free_tag = '[FREE] ' if m.get('free') else ''
                ctx = m.get('context', 0)
                ctx_str = f'  ctx:{ctx}' if ctx else ''
                name = m.get('name', '')
                if name and name != mid:
                    result += f"  {i}. {free_tag}{mid}\n     {name}{ctx_str}\n"
                else:
                    result += f"  {i}. {free_tag}{mid}{ctx_str}\n"

            result += f"\n  Total available: {len(models)} models"
            return result

        except Exception as e:
            return f"Error searching models: {e}\nTry /live_models command or check API key."

    # ═══════════════════════════════════════════════════════════════
    # BROWSER TOOL IMPLEMENTATIONS (v5.4)
    # ═══════════════════════════════════════════════════════════════

    def _bsession(self):
        """Get the shared browser session."""
        from .webcontrol import get_session
        return get_session()

    def _browser_navigate(self, url: str) -> str:
        try:
            result = self._bsession().navigate(url)
            if 'error' in result:
                return f"Error: {result['error']}"
            parts = [
                f"URL: {result['url']}",
                f"Status: {result['status']}",
                f"Title: {result['title']}",
                f"Links: {result['links_total']}",
                f"Forms: {result['forms_total']}",
                f"Images: {result['images_total']}",
            ]
            if result['text']:
                parts.append(f"\n--- Page Content ---\n{result['text']}")
            if result['links']:
                parts.append(f"\n--- Links ({len(result['links'])}) ---")
                for i, link in enumerate(result['links'][:20], 1):
                    parts.append(f"  {i}. {link['text']} -> {link['url']}")
            if result['forms']:
                parts.append(f"\n--- Forms ({len(result['forms'])}) ---")
                for i, form in enumerate(result['forms'][:5], 1):
                    parts.append(f"  {i}. {form['method']} -> {form['action']}")
                    for inp in form['inputs'][:5]:
                        parts.append(f"     - {inp['name']} ({inp['type']})")
            return '\n'.join(parts)
        except Exception as e:
            return f"Error: {e}"

    def _browser_login(self, url: str, username: str, password: str,
                       username_field: str, password_field: str) -> str:
        try:
            result = self._bsession().login(url, username, password,
                                            username_field, password_field)
            if 'error' in result:
                return f"Error: {result['error']}"
            parts = [
                f"URL: {result['url']}",
                f"Status: {result['status']}",
                f"Title: {result['title']}",
                f"Login: {'SUCCESS' if result['success'] else 'FAILED'}",
                f"Redirected: {result.get('redirected', False)}",
            ]
            if not result['success']:
                parts.append("WARNING: Error indicators found on response page.")
            if result.get('text_preview'):
                parts.append(f"\n--- Page Preview ---\n{result['text_preview'][:500]}")
            return '\n'.join(parts)
        except Exception as e:
            return f"Error: {e}"

    def _browser_click(self, target: str, by: str) -> str:
        try:
            result = self._bsession().click(target, by)
            if 'error' in result:
                return f"Error: {result['error']}"
            parts = [
                f"URL: {result['url']}",
                f"Status: {result.get('status', '')}",
                f"Title: {result.get('title', '')}",
            ]
            if result.get('text'):
                text = result['text']
                if len(text) > 3000:
                    text = text[:3000] + '\n... (truncated)'
                parts.append(f"\n--- Page Content ---\n{text}")
            return '\n'.join(parts)
        except Exception as e:
            return f"Error: {e}"

    def _browser_fill_form(self, url: str, form_data_str: str,
                           submit: bool, form_index: int) -> str:
        try:
            form_data = json.loads(form_data_str)
        except json.JSONDecodeError:
            return f"Error: Invalid JSON in form_data: {form_data_str}"
        try:
            result = self._bsession().fill_form(url, form_data, submit, form_index)
            if 'error' in result:
                return f"Error: {result['error']}"
            parts = [
                f"URL: {result['url']}",
                f"Status: {result.get('status', '')}",
                f"Title: {result.get('title', '')}",
            ]
            if result.get('submitted_data'):
                parts.append(f"Submitted: {json.dumps(result['submitted_data'], ensure_ascii=False)}")
            elif result.get('filled_data'):
                parts.append(f"Filled: {json.dumps(result['filled_data'], ensure_ascii=False)}")
                parts.append(result.get('message', ''))
            return '\n'.join(parts)
        except Exception as e:
            return f"Error: {e}"

    def _browser_snapshot(self) -> str:
        try:
            result = self._bsession().snapshot()
            if 'error' in result:
                return f"Error: {result['error']}"
            parts = [
                f"URL: {result['url']}",
                f"Title: {result['title']}",
                f"Canonical: {result.get('canonical', '')}",
            ]
            if result.get('meta'):
                parts.append(f"\n--- Meta ---")
                for k, v in list(result['meta'].items())[:10]:
                    parts.append(f"  {k}: {v}")
            if result.get('headings'):
                parts.append(f"\n--- Headings ({len(result['headings'])}) ---")
                for h in result['headings'][:20]:
                    parts.append(f"  {h['level']}: {h['text']}")
            parts.append(f"\n--- Stats ---")
            parts.append(f"  Links: {result['links_count']} across {len(result.get('links_domains', []))} domains")
            parts.append(f"  Forms: {result['forms_count']}")
            parts.append(f"  Images: {result['images_count']}")
            parts.append(f"  Text length: {result['visible_text_length']} chars")
            parts.append(f"  History: {result['history_length']} pages")
            if result.get('visible_text'):
                text = result['visible_text']
                if len(text) > 5000:
                    text = text[:5000] + '\n... (truncated)'
                parts.append(f"\n--- Visible Text ---\n{text}")
            if result.get('links'):
                parts.append(f"\n--- Top Links ---")
                for link in result['links'][:25]:
                    parts.append(f"  {link['text'][:60]} -> {link['url']}")
            if result.get('forms'):
                parts.append(f"\n--- Forms ---")
                for i, form in enumerate(result['forms'][:5], 1):
                    parts.append(f"  {i}. {form['method']} -> {form['action']}")
                    for field in form['fields'][:8]:
                        req = ' [required]' if field.get('required') else ''
                        parts.append(f"     {field['name']} ({field['type']}){req}")
            return '\n'.join(parts)
        except Exception as e:
            return f"Error: {e}"

    def _browser_extract(self, css_selector: str) -> str:
        try:
            result = self._bsession().extract(css_selector)
            if 'error' in result:
                return f"Error: {result['error']}"
            parts = [f"Found {result['count']} elements for '{css_selector}':"]
            for i, elem in enumerate(result['elements'][:15], 1):
                parts.append(f"\n--- Element {i} ({elem['tag']}) ---")
                parts.append(elem['text'][:1500])
                if elem.get('links'):
                    for link in elem['links'][:5]:
                        parts.append(f"  [LINK] {link['text']} -> {link['url']}")
            return '\n'.join(parts)
        except Exception as e:
            return f"Error: {e}"

    def _browser_download(self, url: str, save_path: str) -> str:
        try:
            result = self._bsession().download(url, save_path)
            if 'error' in result:
                return f"Error: {result['error']}"
            return (f"Downloaded: {result['url']}\n"
                    f"Saved to: {result['saved_to']}\n"
                    f"Size: {result['size_human']} ({result['size_bytes']} bytes)")
        except Exception as e:
            return f"Error: {e}"

    def _browser_action(self, args: dict) -> str:
        action = args.get('action', '')
        try:
            session = self._bsession()
            if action == 'navigate':
                url = args.get('url', '')
                if not url:
                    return "Error: 'url' required for navigate action"
                return self._browser_navigate(url)
            elif action == 'click':
                target = args.get('target', '')
                by = args.get('by', args.get('css_selector', 'text'))
                if by and by.startswith(('.', '#', '[', 'div', 'span', 'a', 'button')):
                    by = 'css'
                if not target:
                    return "Error: 'target' required for click action"
                return self._browser_click(target, by)
            elif action == 'fill':
                url = args.get('url', '')
                data = args.get('data', '{}')
                if not url:
                    return "Error: 'url' required for fill action"
                return self._browser_fill_form(url, data, True, 0)
            elif action == 'submit':
                if not session.soup:
                    return "Error: No page loaded. Navigate first."
                forms = session.soup.find_all('form')
                if forms:
                    result = session._submit_form(forms[0])
                    if 'error' in result:
                        return f"Error: {result['error']}"
                    return f"Submitted form -> {result.get('url', '')} (status {result.get('status', '')})"
                return "Error: No forms found on current page"
            elif action == 'extract':
                selector = args.get('css_selector', args.get('target', ''))
                if not selector:
                    return "Error: 'css_selector' or 'target' required for extract action"
                return self._browser_extract(selector)
            elif action == 'back':
                history = session.history
                if len(history) > 1:
                    prev_url = history[-2]
                    return self._browser_navigate(prev_url)
                return "Error: No previous page in history"
            else:
                return f"Error: Unknown action '{action}'. Use: navigate, click, fill, submit, extract, back"
        except Exception as e:
            return f"Error: {e}"

    def _browser_screenshot(self) -> str:
        try:
            result = self._bsession().screenshot()
            if 'error' in result:
                return f"Error: {result['error']}"
            return result.get('rendering', 'No rendering available')
        except Exception as e:
            return f"Error: {e}"

    def _browser_cookies(self, action: str) -> str:
        try:
            session = self._bsession()
            if action == 'clear':
                return session.clear_cookies()
            result = session.get_cookies()
            if not result['cookies']:
                return "No cookies in session."
            parts = [f"Cookies ({result['count']}):"]
            for name, value in sorted(result['cookies'].items()):
                display_val = value[:50] + '...' if len(value) > 50 else value
                parts.append(f"  {name} = {display_val}")
            return '\n'.join(parts)
        except Exception as e:
            return f"Error: {e}"

    # ═══════════════════════════════════════════════════════════════
    # SELENIUM BROWSER TOOLS (v7.0)
    # Real browser automation using Selenium + Firefox + Geckodriver
    # 12 advanced tools: Navigate, Smart Login, Click, Type, Screenshot,
    # JS Execute, Wait, Scrape, Dropdown, Scroll, Cookies, Close
    # Graceful fallback — tools show install instructions if Selenium unavailable
    # ═══════════════════════════════════════════════════════════════

    def _se_not_available(self) -> str:
        return ('[SELENIUM] Selenium browser is not available or failed to connect.\n'
                'Firefox or geckodriver is not installed/cannot start on this system.\n'
                '\n'
                'INSTEAD, use these HTTP-based browser tools (NO Firefox needed):\n'
                '  - browser_navigate: Open URL and get page info\n'
                '  - browser_snapshot: Full page snapshot (text, links, forms, images)\n'
                '  - browser_click: Click links/buttons by text or CSS selector\n'
                '  - browser_fill_form: Fill and submit forms\n'
                '  - browser_extract: Extract content by CSS selector\n'
                '  - browser_download: Download files from URLs\n'
                '  - browser_action: General web action (navigate, click, fill, etc.)\n'
                '  - browser_screenshot: Visual text-based page rendering\n'
                '\n'
                'These tools use HTTP requests and work on ANY system (no browser install needed).')

    def _se_connection_error(self, error: Exception) -> str:
        """Handle Selenium connection/driver errors with helpful fallback message."""
        err_str = str(error)
        # Check common error patterns
        if 'Connection refused' in err_str or 'Max retries exceeded' in err_str:
            return ('[SELENIUM] Browser connection failed — geckodriver is not running.\n'
                    'Error: ' + err_str[:200] + '\n'
                    '\n'
                    'FIX: Use HTTP-based browser tools instead (no Firefox/geckodriver needed):\n'
                    '  browser_navigate, browser_snapshot, browser_click, browser_fill_form,\n'
                    '  browser_extract, browser_download, browser_action, browser_screenshot')
        elif 'Firefox' in err_str or 'firefox' in err_str:
            return ('[SELENIUM] Firefox browser not found or cannot start.\n'
                    'Error: ' + err_str[:200] + '\n'
                    '\n'
                    'FIX: Use browser_navigate, browser_snapshot, and other browser_* tools instead.\n'
                    'These work via HTTP and need NO browser installation.')
        else:
            return ('[SELENIUM] Browser error: ' + err_str[:300] + '\n'
                    '\n'
                    'FIX: Use browser_navigate, browser_snapshot, and other browser_* tools instead.')

    def _register_selenium_tools(self):
        # Check availability (don't register if missing)
        try:
            from .selenium_browser import SELENIUM_AVAILABLE
            if not SELENIUM_AVAILABLE:
                return
        except ImportError:
            return

        self.register(
            'se_navigate',
            'SELENIUM BROWSER: Open a URL in a real Firefox browser. '
            'Renders full page with JavaScript, returns title, text, link/form/image counts. '
            'Use this as the first step for any Selenium-based interaction. '
            'Supports dynamic content, SPAs, and JavaScript-rendered pages.',
            {
                'type': 'object',
                'properties': {
                    'url': {'type': 'string', 'description': 'Full URL to navigate to'},
                    'wait_seconds': {'type': 'number', 'description': 'Seconds to wait for page load (default 3)'}
                },
                'required': ['url']
            },
            lambda args: self._se_navigate(args['url'], args.get('wait_seconds', 3))
        )

        self.register(
            'se_login',
            'SELENIUM BROWSER: Advanced login automation with smart field detection. '
            'Auto-detects username/password fields by common patterns (name, id, type, placeholder). '
            'Handles form submission, waits for redirect, detects login errors. '
            'Supports custom field names and submit button text. '
            'Maintains cookies after login for subsequent requests.',
            {
                'type': 'object',
                'properties': {
                    'url': {'type': 'string', 'description': 'Login page URL'},
                    'username': {'type': 'string', 'description': 'Username or email'},
                    'password': {'type': 'string', 'description': 'Password'},
                    'username_field': {'type': 'string', 'description': 'Custom CSS/name/id for username field (auto-detect if empty)'},
                    'password_field': {'type': 'string', 'description': 'Custom CSS/name/id for password field (auto-detect if empty)'},
                    'submit_text': {'type': 'string', 'description': 'Text of the submit button to click (auto-detect if empty)'},
                    'wait_for_url': {'type': 'string', 'description': 'Wait until URL contains this string (confirms login success)'},
                    'wait_seconds': {'type': 'number', 'description': 'Seconds to wait after submit (default 5)'}
                },
                'required': ['url', 'username', 'password']
            },
            lambda args: self._se_login(
                args['url'], args['username'], args['password'],
                args.get('username_field', ''), args.get('password_field', ''),
                args.get('submit_text', ''), args.get('wait_for_url', ''),
                args.get('wait_seconds', 5)
            )
        )

        self.register(
            'se_click',
            'SELENIUM BROWSER: Click an element on the current page. '
            'Scrolls into view, waits for clickable, then clicks. '
            'Supports text matching, CSS selector, XPath, ID, or name.',
            {
                'type': 'object',
                'properties': {
                    'target': {'type': 'string', 'description': 'Element text, CSS selector, XPath, ID, or name'},
                    'by': {'type': 'string', 'description': 'Search method: text, css, xpath, id, name (default: text)'},
                    'wait_seconds': {'type': 'number', 'description': 'Max wait for element to be clickable (default 3)'}
                },
                'required': ['target']
            },
            lambda args: self._se_click(args['target'], args.get('by', 'text'), args.get('wait_seconds', 3))
        )

        self.register(
            'se_type',
            'SELENIUM BROWSER: Type text into an input field. '
            'Clears existing text first, types character by character for realism. '
            'Optionally presses Enter after typing.',
            {
                'type': 'object',
                'properties': {
                    'target': {'type': 'string', 'description': 'CSS selector, XPath, ID, or name of the input field'},
                    'text': {'type': 'string', 'description': 'Text to type into the field'},
                    'by': {'type': 'string', 'description': 'Search method: css, xpath, id, name (default: css)'},
                    'clear_first': {'type': 'boolean', 'description': 'Clear field before typing (default true)'},
                    'press_enter': {'type': 'boolean', 'description': 'Press Enter after typing (default false)'}
                },
                'required': ['target', 'text']
            },
            lambda args: self._se_type(
                args['target'], args['text'], args.get('by', 'css'),
                args.get('clear_first', True), args.get('press_enter', False)
            )
        )

        self.register(
            'se_screenshot',
            'SELENIUM BROWSER: Take a real PNG screenshot of the current page. '
            'Saves to temp directory. Supports full-page scrolling screenshots. '
            'Returns file path and size.',
            {
                'type': 'object',
                'properties': {
                    'filename': {'type': 'string', 'description': 'Output filename (auto-generated if empty)'},
                    'full_page': {'type': 'boolean', 'description': 'Capture full scrolling page (default false)'}
                },
                'required': []
            },
            lambda args: self._se_screenshot(args.get('filename', ''), args.get('full_page', False))
        )

        self.register(
            'se_execute_js',
            'SELENIUM BROWSER: Execute JavaScript code in the browser page. '
            'Returns the result. Use for custom interactions, data extraction, '
            'or page manipulation. Has access to full DOM API.',
            {
                'type': 'object',
                'properties': {
                    'script': {'type': 'string', 'description': 'JavaScript code to execute'}
                },
                'required': ['script']
            },
            lambda args: self._se_execute_js(args['script'])
        )

        self.register(
            'se_wait_for',
            'SELENIUM BROWSER: Smart wait for a condition to become true. '
            'Conditions: element, visible, clickable, text, url_contains, title_contains, page_load, alert, gone.',
            {
                'type': 'object',
                'properties': {
                    'condition': {'type': 'string', 'description': 'Wait condition: element, visible, clickable, text, url_contains, title_contains, page_load, alert, gone'},
                    'value': {'type': 'string', 'description': 'Value to wait for (CSS selector, text, URL, etc.)'},
                    'timeout': {'type': 'number', 'description': 'Max wait in seconds (default 10)'}
                },
                'required': ['condition', 'value']
            },
            lambda args: self._se_wait_for(args['condition'], args['value'], args.get('timeout', 10))
        )

        self.register(
            'se_scrape',
            'SELENIUM BROWSER: Scrape content from the current page. '
            'Extract types: text, links, images, forms, tables, html, meta, all. '
            'Optionally filter by CSS selector.',
            {
                'type': 'object',
                'properties': {
                    'extract': {'type': 'string', 'description': 'What to extract: text, links, images, forms, tables, html, meta, all (default: text)'},
                    'css_selector': {'type': 'string', 'description': 'Optional CSS selector to filter elements'},
                    'max_items': {'type': 'integer', 'description': 'Max items to return (default 50)'}
                },
                'required': []
            },
            lambda args: self._se_scrape(args.get('extract', 'text'), args.get('css_selector', ''), args.get('max_items', 50))
        )

        self.register(
            'se_dropdown',
            'SELENIUM BROWSER: Select an option from a dropdown (select) element.',
            {
                'type': 'object',
                'properties': {
                    'target': {'type': 'string', 'description': 'CSS selector, XPath, ID, or name of the select element'},
                    'value': {'type': 'string', 'description': 'Option to select (text, value, or index number)'},
                    'by': {'type': 'string', 'description': 'Search method for dropdown: css, xpath, id, name (default: css)'},
                    'select_by': {'type': 'string', 'description': 'How to select: value, text, index (default: value)'}
                },
                'required': ['target', 'value']
            },
            lambda args: self._se_dropdown(args['target'], args['value'], args.get('by', 'css'), args.get('select_by', 'value'))
        )

        self.register(
            'se_scroll',
            'SELENIUM BROWSER: Scroll the page in any direction or to an element. '
            'Directions: down, up, top, bottom. Can also scroll to a specific element.',
            {
                'type': 'object',
                'properties': {
                    'direction': {'type': 'string', 'description': 'Scroll direction: down, up, top, bottom (default: down)'},
                    'amount': {'type': 'integer', 'description': 'Pixels to scroll (default 500)'},
                    'to_element': {'type': 'string', 'description': 'CSS selector of element to scroll into view'},
                    'by': {'type': 'string', 'description': 'Search method for to_element: css, xpath (default: css)'}
                },
                'required': []
            },
            lambda args: self._se_scroll(
                args.get('direction', 'down'), args.get('amount', 500),
                args.get('to_element', ''), args.get('by', 'css')
            )
        )

        self.register(
            'se_cookies',
            'SELENIUM BROWSER: View or manage browser cookies. '
            'Actions: view, clear, clear_storage (localStorage + sessionStorage).',
            {
                'type': 'object',
                'properties': {
                    'action': {'type': 'string', 'description': 'Action: view (default), clear, clear_storage'},
                    'domain': {'type': 'string', 'description': 'Filter cookies by domain (view action only)'}
                },
                'required': []
            },
            lambda args: self._se_cookies(args.get('action', 'view'), args.get('domain', ''))
        )

        self.register(
            'se_close',
            'SELENIUM BROWSER: Close the browser session and free resources. '
            'Call this when done with browser automation to release memory.',
            {
                'type': 'object',
                'properties': {},
                'required': []
            },
            lambda args: self._se_close()
        )

        # ══════════════════════════════════════════════════
        # ADVANCED AUTH AUTOMATION TOOLS (v7.3)
        # Google OAuth, Auth Detection, 2FA, CAPTCHA,
        # Popup Windows, GUI Mode, Cookie Import/Export
        # ══════════════════════════════════════════════════

        self.register(
            'se_google_login',
            'SELENIUM BROWSER: Full Google "Sign in with Google" automation. '
            'Handles: click Google button, email input, password input, account picker, '
            '2FA/OTP code, consent/permission screen, CAPTCHA detection. '
            'If CAPTCHA detected, returns needs_gui=true for manual fallback.',
            {
                'type': 'object',
                'properties': {
                    'url': {'type': 'string', 'description': 'URL of the page with Google login button'},
                    'email': {'type': 'string', 'description': 'Google email/phone'},
                    'password': {'type': 'string', 'description': 'Google password'},
                    'otp_code': {'type': 'string', 'description': '2FA/OTP code if account has 2-step verification (optional)'},
                    'button_text': {'type': 'string', 'description': 'Text of the "Sign in with Google" button (auto-detect if empty)'}
                },
                'required': ['url', 'email', 'password']
            },
            lambda args: self._se_google_login(
                args['url'], args['email'], args['password'],
                args.get('otp_code', ''), args.get('button_text', '')
            )
        )

        self.register(
            'se_detect_auth',
            'SELENIUM BROWSER: Detect what type of authentication form is on the current page. '
            'Identifies: Google OAuth, GitHub OAuth, Facebook OAuth, Twitter/X OAuth, '
            'Discord OAuth, Microsoft OAuth, Apple OAuth, basic login forms, 2FA, CAPTCHA, SSO/SAML. '
            'Returns detected fields and suggested tools to use.',
            {
                'type': 'object',
                'properties': {},
                'required': []
            },
            lambda args: self._se_detect_auth()
        )

        self.register(
            'se_open_gui',
            'SELENIUM BROWSER: Open GUI mode when automated login fails. '
            'On Termux: launches Termux X11/VNC and opens the URL for manual auth completion. '
            'On Desktop: opens the URL in the system browser for manual login. '
            'After manual completion, use se_import_cookies to bring cookies back.',
            {
                'type': 'object',
                'properties': {
                    'url': {'type': 'string', 'description': 'URL to open in GUI browser (default: current page)'},
                    'purpose': {'type': 'string', 'description': 'Purpose description (default: "login")'}
                },
                'required': []
            },
            lambda args: self._se_open_gui(args.get('url', ''), args.get('purpose', 'login'))
        )

        self.register(
            'se_popup',
            'SELENIUM BROWSER: Manage popup windows and tabs opened by OAuth/social login. '
            'Actions: list (all open windows), switch (to specific window), close_popup (close all popups). '
            'Use after clicking social login buttons that open new tabs.',
            {
                'type': 'object',
                'properties': {
                    'action': {'type': 'string', 'description': 'Action: list, switch, close_popup (default: list)'},
                    'window_index': {'type': 'integer', 'description': 'Window index to switch to (for switch action, default 1)'}
                },
                'required': []
            },
            lambda args: self._se_popup(args.get('action', 'list'), args.get('window_index', 1))
        )

        self.register(
            'se_switch_frame',
            'SELENIUM BROWSER: Switch to or list iframes. '
            'Many OAuth login buttons are embedded in iframes. '
            'Use list action to see all iframes, then switch to interact with them. '
            'Use "main" to switch back to main content.',
            {
                'type': 'object',
                'properties': {
                    'frame': {'type': 'string', 'description': 'Frame index, id, name, or "main" (default: empty = list all)'}
                },
                'required': []
            },
            lambda args: self._se_switch_frame(args.get('frame', ''))
        )

        self.register(
            'se_import_cookies',
            'SELENIUM BROWSER: Import cookies from JSON file or JSON string. '
            'Use after manual login via se_open_gui to import cookies back. '
            'Restores login session without re-entering credentials.',
            {
                'type': 'object',
                'properties': {
                    'source': {'type': 'string', 'description': 'JSON file path or JSON string containing cookies array'},
                    'is_file': {'type': 'boolean', 'description': 'True if source is a file path, False if JSON string (default true)'}
                },
                'required': ['source']
            },
            lambda args: self._se_import_cookies(args['source'], args.get('is_file', True))
        )

        self.register(
            'se_export_cookies',
            'SELENIUM BROWSER: Export current browser cookies to a JSON file. '
            'Save cookies after login so you can import them later. '
            'Useful for persisting login sessions across restarts.',
            {
                'type': 'object',
                'properties': {
                    'save_to': {'type': 'string', 'description': 'File path to save cookies JSON (default: auto-generated)'}
                },
                'required': []
            },
            lambda args: self._se_export_cookies(args.get('save_to', ''))
        )

        self.register(
            'se_upload',
            'SELENIUM BROWSER: Upload a file to a file input element on the page.',
            {
                'type': 'object',
                'properties': {
                    'target': {'type': 'string', 'description': 'CSS selector of the file input element'},
                    'file_path': {'type': 'string', 'description': 'Path to the file to upload'}
                },
                'required': ['target', 'file_path']
            },
            lambda args: self._se_upload(args['target'], args['file_path'])
        )

        # ── v7.4: Auth Automation Tools ──

        self.register(
            'se_auth_google',
            'AUTH: Automate Google OAuth login. Handles email/password fields, 2FA detection, '
            'iframe switching, and popup handling. If login fails, suggests GUI mode switch. '
            'IMPORTANT: For 2FA/CAPTCHA, use se_switch_to_gui first for manual completion.',
            {
                'type': 'object',
                'properties': {
                    'email': {'type': 'string', 'description': 'Google account email'},
                    'password': {'type': 'string', 'description': 'Google account password'},
                    'login_url': {'type': 'string', 'description': 'Login URL (default: Google sign-in page)'},
                    'timeout': {'type': 'number', 'description': 'Max wait time in seconds (default 60)'},
                },
                'required': ['email'],
            },
            lambda args: self._se_get_session().auth_google(
                args['email'], args.get('password', ''),
                args.get('login_url', ''), args.get('timeout', 60)
            )
        )

        self.register(
            'se_auth_generic',
            'AUTH: Generic login automation with GUI fallback. Tries headless first, '
            'auto-switches to GUI browser if login fails. Good for complex auth flows, '
            'CAPTCHAs, or sites that block headless browsers.',
            {
                'type': 'object',
                'properties': {
                    'url': {'type': 'string', 'description': 'Login page URL'},
                    'email': {'type': 'string', 'description': 'Email/username'},
                    'password': {'type': 'string', 'description': 'Password'},
                    'email_selector': {'type': 'string', 'description': 'CSS selector for email field (auto-detect if empty)'},
                    'password_selector': {'type': 'string', 'description': 'CSS selector for password field (auto-detect if empty)'},
                    'submit_selector': {'type': 'string', 'description': 'CSS selector for submit button (auto-detect if empty)'},
                    'wait_for_url': {'type': 'string', 'description': 'Wait for URL containing this string after login'},
                    'auto_gui_fallback': {'type': 'boolean', 'description': 'Auto-switch to GUI if headless fails (default true)'},
                },
                'required': ['url'],
            },
            lambda args: self._se_get_session().auth_generic(
                args['url'], args.get('email', ''), args.get('password', ''),
                args.get('email_selector', ''), args.get('password_selector', ''),
                args.get('submit_selector', ''), args.get('wait_for_url', ''),
                args.get('auto_gui_fallback', True),
            )
        )

        self.register(
            'se_switch_to_gui',
            'AUTH: Switch browser from headless to GUI mode. On Termux: auto-launches Termux X11. '
            'On Desktop: uses native X11 display. Required for completing logins with CAPTCHA, '
            '2FA, or other interactive challenges that need a visible browser.',
            {
                'type': 'object',
                'properties': {
                    'force': {'type': 'boolean', 'description': 'Force re-initialize even if already in GUI mode (default false)'},
                },
                'required': [],
            },
            lambda args: self._se_get_session().switch_to_gui_mode(force=args.get('force', False))
        )

        self.register(
            'se_switch_to_headless',
            'AUTH: Switch browser back to headless mode (no visible window).',
            {
                'type': 'object',
                'properties': {},
                'required': [],
            },
            lambda args: self._se_get_session().switch_to_headless()
        )

        self.register(
            'se_handle_popup',
            'AUTH: Handle browser alerts/confirms/prompts. Accept or dismiss popups.',
            {
                'type': 'object',
                'properties': {
                    'action': {'type': 'string', 'description': 'Action: accept, dismiss'},
                    'timeout': {'type': 'number', 'description': 'Wait timeout in seconds (default 5)'},
                },
                'required': [],
            },
            lambda args: self._se_get_session().handle_popup(
                args.get('action', 'accept'), args.get('timeout', 5)
            )
        )

        self.register(
            'se_switch_tab',
            'AUTH: Switch between browser tabs/windows. Google OAuth often opens new tabs.',
            {
                'type': 'object',
                'properties': {
                    'tab_index': {'type': 'integer', 'description': 'Tab index (0=first, -1=newest, default -1)'},
                },
                'required': [],
            },
            lambda args: self._se_get_session().switch_tab(tab_index=args.get('tab_index', -1))
        )

        self.register(
            'se_close_tab',
            'AUTH: Close a browser tab (cannot close the last tab).',
            {
                'type': 'object',
                'properties': {
                    'tab_index': {'type': 'integer', 'description': 'Tab index to close (default -1 = newest)'},
                },
                'required': [],
            },
            lambda args: self._se_get_session().close_tab(tab_index=args.get('tab_index', -1))
        )

        # ── Skill Management Tools ──
        self.register(
            'list_skills',
            'List all installed agent skills with descriptions.',
            {
                'type': 'object',
                'properties': {},
                'required': [],
            },
            lambda args: self._list_skills()
        )

        self.register(
            'read_skill',
            'Read the content of an installed skill (SKILL.md). Use list_skills first to see available skills.',
            {
                'type': 'object',
                'properties': {
                    'name': {'type': 'string', 'description': 'Skill name (e.g. "find-skills", "pdf")'},
                },
                'required': ['name'],
            },
            lambda args: self._read_skill(args['name'])
        )

    # ── Selenium Tool Handler Implementations ───────────────

    def _se_get_session(self):
        """Get the Selenium browser session."""
        try:
            from .selenium_browser import get_selenium_session
            return get_selenium_session(headless=True)
        except ImportError:
            return None

    def _se_navigate(self, url: str, wait_seconds: float) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            result = session.navigate(url, wait_seconds)
            if 'error' in result:
                return f"[ERROR] {result['error']}"
            return json.dumps(result, indent=2, ensure_ascii=False, default=str)
        except Exception as e:
            return self._se_connection_error(e)

    def _se_login(self, url: str, username: str, password: str,
                  username_field: str, password_field: str,
                  submit_text: str, wait_for_url: str, wait_seconds: float) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            result = session.smart_login(
                url, username, password,
                username_field, password_field,
                submit_text, wait_for_url, wait_seconds
            )
            if 'error' in result and not result.get('success'):
                return f"[ERROR] Login failed: {result['error']}\nSteps:\n" + "\n".join(f"  - {s}" for s in result.get('steps', []))
            output = [
                f"Login: {'SUCCESS' if result.get('success') else 'FAILED'}",
                f"URL: {result.get('url', '')}",
                f"Title: {result.get('title', '')}",
                f"Cookies: {result.get('cookies', 'N/A')}",
                "Steps:"
            ]
            for s in result.get('steps', []):
                output.append(f"  - {s}")
            return '\n'.join(output)
        except Exception as e:
            return self._se_connection_error(e)

    def _se_click(self, target: str, by: str, wait_seconds: float) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            result = session.click_element(target, by, wait_seconds)
            if 'error' in result:
                return f"[ERROR] {result['error']}"
            return json.dumps(result, indent=2, ensure_ascii=False, default=str)
        except Exception as e:
            return self._se_connection_error(e)

    def _se_type(self, target: str, text: str, by: str, clear_first: bool, press_enter: bool) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            result = session.type_text(target, text, by, clear_first, press_enter)
            if 'error' in result:
                return f"[ERROR] {result['error']}"
            return json.dumps(result, indent=2, ensure_ascii=False, default=str)
        except Exception as e:
            return self._se_connection_error(e)

    def _se_screenshot(self, filename: str, full_page: bool) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            result = session.take_screenshot(filename, full_page)
            if 'error' in result:
                return f"[ERROR] {result['error']}"
            return (f"Screenshot saved: {result['saved_to']}\n"
                    f"Size: {result['size_human']}\n"
                    f"Full page: {result['full_page']}")
        except Exception as e:
            return self._se_connection_error(e)

    def _se_execute_js(self, script: str) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            result = session.execute_js(script)
            if 'error' in result:
                return f"[ERROR] {result['error']}"
            return f"Result ({result['result_type']}): {result['result']}"
        except Exception as e:
            return self._se_connection_error(e)

    def _se_wait_for(self, condition: str, value: str, timeout: float) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            result = session.wait_for(condition, value, timeout)
            if 'error' in result:
                return f"[ERROR] {result['error']}"
            return f"Wait satisfied: {result['condition']}"
        except Exception as e:
            return self._se_connection_error(e)

    def _se_scrape(self, extract: str, css_selector: str, max_items: int) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            result = session.scrape_page(extract, css_selector, max_items)
            if 'error' in result:
                return f"[ERROR] {result['error']}"
            return json.dumps(result, indent=2, ensure_ascii=False, default=str)
        except Exception as e:
            return self._se_connection_error(e)

    def _se_dropdown(self, target: str, value: str, by: str, select_by: str) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            result = session.select_dropdown(target, value, by, select_by)
            if 'error' in result:
                return f"[ERROR] {result['error']}"
            return json.dumps(result, indent=2, ensure_ascii=False, default=str)
        except Exception as e:
            return self._se_connection_error(e)

    def _se_scroll(self, direction: str, amount: int, to_element: str, by: str) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            result = session.scroll_page(direction, amount, to_element, by)
            if 'error' in result:
                return f"[ERROR] {result['error']}"
            return json.dumps(result, indent=2, ensure_ascii=False, default=str)
        except Exception as e:
            return self._se_connection_error(e)

    def _se_cookies(self, action: str, domain: str) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            if action == 'clear':
                result = session.clear_cookies()
                return json.dumps(result, indent=2, ensure_ascii=False, default=str)
            elif action == 'clear_storage':
                result = session.clear_storage()
                return json.dumps(result, indent=2, ensure_ascii=False, default=str)
            else:
                result = session.get_cookies(domain)
                if 'error' in result:
                    return f"[ERROR] {result['error']}"
                output = [f"Browser Cookies ({result['count']} total)"]
                output.append(f"Domains: {', '.join(result['domains'][:10])}")
                for c in result['cookies'][:20]:
                    display_val = str(c.get('value', ''))[:30] + '...' if len(str(c.get('value', ''))) > 30 else str(c.get('value', ''))
                    output.append(f"  {c.get('name', '?')} = {display_val} (domain: {c.get('domain', '?')})")
                if len(result['cookies']) > 20:
                    output.append(f"  ... and {len(result['cookies']) - 20} more")
                return '\n'.join(output)
        except Exception as e:
            return self._se_connection_error(e)

    def _se_close(self) -> str:
        try:
            from .selenium_browser import close_selenium_session
            close_selenium_session()
            return "Selenium browser session closed."
        except ImportError:
            return self._se_not_available()

    # ── Advanced Auth Tool Handlers (v7.3) ────────────────

    def _se_google_login(self, url: str, email: str, password: str,
                         otp_code: str, button_text: str) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            result = session.google_oauth_login(url, email, password, otp_code, button_text)
            # Format output
            lines = []
            success = result.get('success', False)
            needs_gui = result.get('needs_gui', False)
            needs_otp = result.get('needs_otp', False)

            if success:
                lines.append('[OK] Google Login SUCCESSFUL')
                lines.append(f'URL: {result.get("url", "")}')
                lines.append(f'Title: {result.get("title", "")}')
                lines.append(f'Cookies: {result.get("cookies", "N/A")}')
            elif needs_gui:
                lines.append('[CAPTCHA/GUI] Automated login blocked — manual intervention needed')
                captcha = result.get('captcha', {})
                if captcha.get('found'):
                    lines.append(f'CAPTCHA type: {captcha.get("type", "unknown")}')
                lines.append('Run se_open_gui to open in GUI browser for manual completion')
            elif needs_otp:
                lines.append('[2FA] Two-factor authentication required')
                lines.append('Call se_google_login again with otp_code parameter')
            else:
                lines.append(f'[FAILED] Google login failed')
                if result.get('message'):
                    lines.append(f'Reason: {result["message"]}')
                if result.get('error'):
                    lines.append(f'Error: {result["error"]}')

            for s in result.get('steps', []):
                lines.append(f'  {s}')

            return '\n'.join(lines)
        except Exception as e:
            return self._se_connection_error(e)

    def _se_detect_auth(self) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            result = session.detect_auth_form()
            lines = []
            lines.append(f'URL: {result.get("url", "")}')
            lines.append(f'Auth Type: {result.get("auth_type", "unknown")}')
            lines.append(f'Fields: {", ".join(result.get("fields", [])) or "none"}')

            captcha = result.get('has_captcha', False)
            if captcha:
                lines.append(f'CAPTCHA: YES ({result.get("captcha_type", "unknown")})')

            for s in result.get('suggestions', []):
                lines.append(f'Tip: {s}')

            if result.get('error'):
                lines.append(f'Error: {result["error"]}')

            return '\n'.join(lines)
        except Exception as e:
            return self._se_connection_error(e)

    def _se_open_gui(self, url: str, purpose: str) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            result = session.launch_gui_mode(url or None, purpose)
            lines = []
            lines.append(f'Environment: {result.get("environment", "unknown")}')
            lines.append(f'Purpose: {result.get("purpose", purpose)}')

            if result.get('url'):
                lines.append(f'URL: {result["url"]}')
            if result.get('opened_url'):
                lines.append(f'Opened: {result["opened_url"]}')
            if result.get('display'):
                lines.append(f'Display: {result["display"]}')
            if result.get('vnc_status'):
                lines.append(f'VNC: {result["vnc_status"]}')
            if result.get('success'):
                lines.append(f'Status: Browser opened')
            if result.get('error'):
                lines.append(f'Error: {result["error"]}')
            if result.get('instructions'):
                for inst_line in result['instructions'].split('\n'):
                    lines.append(inst_line)
            if result.get('suggestions'):
                lines.append('Suggestions:')
                for s in result['suggestions']:
                    lines.append(f'  - {s}')

            return '\n'.join(lines)
        except Exception as e:
            return self._se_connection_error(e)

    def _se_popup(self, action: str, window_index: int) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            result = session.handle_popup(action, window_index)

            if action == 'list':
                windows = result.get('windows', [])
                lines = [f'Open Windows ({len(windows)}):']
                for w in windows:
                    current = ' [CURRENT]' if w.get('is_current') else ''
                    if w.get('error'):
                        lines.append(f'  [{w.get("index")}] Error')
                    else:
                        lines.append(f'  [{w.get("index")}] {w.get("title", "N/A")[:60]}{current}')
                        lines.append(f'       {w.get("url", "")[:80]}')
                return '\n'.join(lines)
            elif action == 'switch':
                return f'Switched to window [{window_index}]\nTitle: {result.get("title", "")}\nURL: {result.get("url", "")}'
            elif action == 'close_popup':
                return f'Closed {result.get("closed", 0)} popup(s)\nRemaining: {result.get("remaining", 1)} window(s)\nURL: {result.get("url", "")}'
            else:
                return f'[ERROR] Unknown action: {action}. Use: list, switch, close_popup'
        except Exception as e:
            return self._se_connection_error(e)

    def _se_switch_frame(self, frame: str) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            if frame == 'main':
                result = session.switch_to_main()
                return f'Switched to main content\nURL: {result.get("url", "")}'
            elif frame:
                result = session.switch_to_frame(frame)
                if 'error' in result:
                    return f'[ERROR] {result["error"]}'
                return f'Switched to frame: {frame}\nURL: {result.get("current_url", "")}'
            else:
                # List iframes
                result = session.switch_to_frame('')
                iframes = result.get('iframes', [])
                if not iframes:
                    return 'No iframes found on current page.'
                lines = [f'Iframes ({len(iframes)}):']
                for f in iframes:
                    if f.get('error'):
                        lines.append(f'  [{f.get("index")}] Error')
                    else:
                        lines.append(f'  [{f.get("index")}] id="{f.get("id", "")}" name="{f.get("name", "")}"')
                        lines.append(f'         src="{f.get("src", "")}" displayed={f.get("displayed", "")}')
                return '\n'.join(lines)
        except Exception as e:
            return self._se_connection_error(e)

    def _se_import_cookies(self, source: str, is_file: bool) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            if is_file:
                p = os.path.expanduser(source)
                if not os.path.exists(p):
                    return f'[ERROR] File not found: {source}'
                with open(p, 'r') as f:
                    cookies = json.load(f)
            else:
                cookies = json.loads(source)

            if not isinstance(cookies, list):
                return f'[ERROR] Cookies must be a JSON array, got {type(cookies).__name__}'

            result = session.set_cookies(cookies)
            if 'error' in result:
                return f'[ERROR] {result["error"]}'
            return f'Imported {result["cookies_set"]} cookie(s) successfully.\nNavigate to the site to verify the login session.'
        except json.JSONDecodeError as e:
            return f'[ERROR] Invalid JSON: {e}'
        except Exception as e:
            return self._se_connection_error(e)

    def _se_export_cookies(self, save_to: str) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            result = session.extract_cookies(save_to)
            if 'error' in result:
                return f'[ERROR] {result["error"]}'
            return (f'Exported {result["count"]} cookie(s)\n'
                    f'Saved to: {result["saved_to"]}\n'
                    f'Import later with: se_import_cookies source="{result["saved_to"]}"')
        except Exception as e:
            return self._se_connection_error(e)

    def _se_upload(self, target: str, file_path: str) -> str:
        session = self._se_get_session()
        if not session:
            return self._se_not_available()
        try:
            result = session.upload_file(target, file_path)
            if 'error' in result:
                return f'[ERROR] {result["error"]}'
            return (f'File uploaded: {os.path.basename(file_path)}\n'
                    f'Element: {result["element"]}\n'
                    f'Full path: {result["file"]}')
        except Exception as e:
            return self._se_connection_error(e)

    # ── Skill Helpers ─────────────────────────────────────

    def _list_skills(self) -> str:
        """List all installed skills with descriptions."""
        skills_dir = os.path.expanduser('~/.agents/skills')
        if not os.path.isdir(skills_dir):
            return 'No skills directory found. Use /install to install skills.'
        names = sorted([d for d in os.listdir(skills_dir) if os.path.isdir(os.path.join(skills_dir, d))])
        if not names:
            return 'No skills installed. Use /install to install skills.'
        result = []
        result.append(f'Installed Skills ({len(names)}):')
        for name in names:
            skill_md = os.path.join(skills_dir, name, 'SKILL.md')
            desc = ''
            if os.path.isfile(skill_md):
                try:
                    with open(skill_md, 'r', encoding='utf-8') as f:
                        first_line = f.readline().strip()
                except Exception:
                    first_line = "No description"
                if first_line.startswith('# '):
                    desc = first_line[2:]
            if desc:
                result.append(f'  {name} \u2014 {desc}')
            else:
                result.append(f'  {name}')
        return '\n'.join(result)

    def _read_skill(self, name: str) -> str:
        """Read the SKILL.md content of a specific skill."""
        skills_dir = os.path.expanduser('~/.agents/skills')
        skill_path = os.path.join(skills_dir, name)
        skill_md = os.path.join(skill_path, 'SKILL.md')
        if not os.path.isdir(skill_path):
            available = [d for d in os.listdir(skills_dir)
                         if os.path.isdir(os.path.join(skills_dir, d))]
            return f'[ERROR] Skill "{name}" not found. Available: {", ".join(sorted(available))}'
        if not os.path.isfile(skill_md):
            return f'[ERROR] No SKILL.md found for skill "{name}".'
        try:
            with open(skill_md, 'r', encoding='utf-8') as f:
                content = f.read()
            return content.strip() or f'[ERROR] SKILL.md for "{name}" is empty.'
        except Exception as e:
            return f'[ERROR] Failed to read skill "{name}": {e}'
